# -*- coding: utf-8 -*-
"""
风味物质自动检索 Web 程序
========================
运行: python3.11 app.py  (默认 http://127.0.0.1:8080)
功能:
  1) 单物质检索: 输入英文/中文物质名 -> 中文名/类别/阈值/介质/备注/匹配方式
  2) 批量检索:   粘贴多行物质名(或上传 txt/csv) -> 表格 + 导出
  3) GC-MS 报告: 上传 Agilent MassHunter 导出的 Excel -> 自动提取并按响应比=含量生成中文+分类+阈值表
"""
import io
import csv
import os
import re
import tempfile
from flask import Flask, request, jsonify, Response, render_template_string

from flavor_core import (match_compound, parse_gcms_excel, enrich, auto_classify,
                         COMPOUNDS, parse_report_text, extract_substance_names,
                         compute_concentration, compute_concentration_list, parse_conc_unit,
                         parse_gcms_pdf_multi,
                         sort_by_completeness, completeness, completeness_label, has_field)
import pdfplumber


app = Flask(__name__)

CAT_COLOR = {
    '醛类': '#E8590C', '酮类': '#B8860B', '酯类': '#2F9E44', '酸类': '#1971C2',
    '醇类': '#5C5C5C', '萜烯类': '#6741D9', '含硫化合物': '#E03131', '吡嗪类': '#0C8599',
    '内酯类': '#C2255C', '其他': '#868E96',
}

PAGE = r"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<meta name="google" content="notranslate">
<meta http-equiv="Content-Language" content="zh-CN, en">
<title>风味物质自动检索 · 分类 · 阈值</title>
<style>
  /* ====== 配色：参考图片精确采样（粉蓝 #7DD3E5 + 珊瑚粉 #F47C6B + 草绿 #5DBE6C + 深蓝灰 #1F2A3D） ====== */
  :root{
    --bg:#e8effa;             /* 页面最外层背景：淡蓝灰 */
    --bg2:#dde7f4;            /* 顶部主内容卡辅底 */
    --card:#ffffff;
    --glass:rgba(255,255,255,.62);      /* 毛玻璃卡面 */
    --glass-line:rgba(255,255,255,.85);
    --card-soft:#eef4fb;      /* 次级卡片/输入框底 */
    --line:#e4ecf5;
    --line-soft:#eef2f7;
    --txt:#1f2a3d;            /* 主文字：深蓝灰（Dashboard 大标题级） */
    --mut:#6b7889;            /* 次要灰（副文案） */
    --mut-2:#a6b2c2;          /* 弱化灰（续文） */
    /* 主色：粉蓝（周图表/Wellness 弧形/图标/Home tab） */
    --teal:#7dd3e5;           /* 粉蓝 */
    --teal-d:#5bc0d4;
    --teal-bg:#e5f3f8;        /* 极淡粉蓝底（Cardio 胶囊等） */
    --teal-ic:#7dd3e5;
    /* 辅色：珊瑚粉（火焰图标/+13% 徽章/Wellness 弧形配色） */
    --coral:#f47c6b;          /* 珊瑚粉 */
    --coral-d:#e9695a;
    --coral-bg:#fdeae7;
    /* 状态色：草绿 / 草绿浅底 */
    --good:#5dbe6c;           /* 草绿（Great level / 血糖 OK） */
    --good-bg:#e8f7ec;
    --warn:#f59e0b;           /* 提醒橙（保留消耗提醒语义） */
    --warn-bg:#fff4e0;
    --grad:linear-gradient(90deg,#5dbed4,#7dd3e5,#f47c6b);  /* 渐变：粉蓝→珊瑚粉 */
    /* 点缀（柔和版保留层次） */
    --sky:#a8d8ef;
    --sky-bg:#e1f1fa;
    --sky-d:#5aa9da;
    --sun:#fcd9a0;
    --sun-bg:#fff3df;
    --rose:#f5a3a3;
    --rose-bg:#fde6e6;
    --lav:#cdb4f0;
    --lav-bg:#f0e8fb;
    --indigo:#b9c4f2;
    --indigo-bg:#eaeefb;
    --rose-2:#f7c2d2;
    --rose-2-bg:#fde7ef;
    --line-teal:#9bd5c4;
    /* 类别徽章（克制：浅底+柔色文字） */
    --c-醛类:#e06b7d;  --cb-醛类:#fde6ec;
    --c-酮类:#d49a4e;  --cb-酮类:#fdf1de;
    --c-酯类:#3f9fd6;  --cb-酯类:#e3f0fb;
    --c-酸类:#5aa9da;  --cb-酸类:#e5f3f8;
    --c-醇类:#7d8ca3;  --cb-醇类:#eef1f6;
    --c-萜烯类与含氧萜类:#9b7fe0; --cb-萜烯类与含氧萜类:#ece4f8;
    --c-含硫化合物:#e07a7a;--cb-含硫化合物:#fbe6e6;
    --c-吡嗪类:#3fafc9;--cb-吡嗪类:#dff0ee;
    --c-内酯类:#d97ba6;--cb-内酯类:#fbe3ed;
    --c-含氮杂环其他含氮化合物:#845ef7;--cb-含氮杂环其他含氮化合物:#f3f0ff;
    --c-呋喃呋喃酮类:#f76707;--cb-呋喃呋喃酮类:#fff0d9;
    --c-噻唑类噻唑啉类:#40c057;--cb-噻唑类噻唑啉类:#e6fcf5;
    --c-其他:#7d8ca3;  --cb-其他:#eef1f6;
    --shadow-sm:0 1px 2px rgba(31,42,61,.05);
    --shadow:0 4px 16px rgba(31,42,61,.07);
    --shadow-lg:0 12px 32px rgba(125,211,229,.12);
    --radius:18px;
    --radius-sm:12px;
  }
  *{box-sizing:border-box}
  body{margin:0;font-family:Arial,"Helvetica Neue","PingFang SC","Microsoft YaHei",sans-serif;background:radial-gradient(1200px 600px at 80% -10%,#e8f0ff 0%,transparent 60%),radial-gradient(900px 500px at -10% 10%,#fdece9 0%,transparent 55%),var(--bg);color:var(--txt);padding:32px 20px 64px;line-height:1.5;-webkit-font-smoothing:antialiased;-moz-osx-font-smoothing:grayscale;font-feature-settings:"tnum" 1,"ss01" 1;letter-spacing:-.005em}
  .page{max-width:1180px;margin:0 auto;}
  h1{font-size:34px;font-weight:800;letter-spacing:-.025em;margin:0 0 6px;color:var(--txt);display:flex;align-items:center;gap:14px;font-variant-numeric:tabular-nums}
  h1 .logo{width:36px;height:36px;border-radius:12px;background:linear-gradient(135deg,var(--teal),var(--coral));color:#fff;display:inline-flex;align-items:center;justify-content:center;font-size:18px;font-weight:800;box-shadow:0 4px 12px rgba(125,211,229,.4)}
  .sub{color:var(--mut);font-size:14px;margin:0 0 22px}
  .sub b{color:var(--teal-d);font-weight:700}

  /* ===== 顶部导航（毛玻璃胶囊，active 信任蓝实心） ===== */
  .tabs{display:flex;gap:6px;flex-wrap:wrap;background:var(--glass);backdrop-filter:blur(12px);-webkit-backdrop-filter:blur(12px);padding:6px;border-radius:16px;box-shadow:var(--shadow);margin-bottom:22px;border:1px solid var(--glass-line)}
  .tab{padding:10px 20px;border-radius:12px;cursor:pointer;background:transparent;color:var(--mut);font-size:14px;font-weight:600;transition:all .2s ease;display:flex;align-items:center;gap:6px}
  .tab:hover{color:var(--teal-d);background:var(--teal-bg)}
  .tab.active{background:var(--teal);color:#1f2a3d;box-shadow:0 4px 12px rgba(125,211,229,.4);font-weight:700}

  /* ===== 卡片（毛玻璃质感） ===== */
  .card{background:var(--glass);backdrop-filter:blur(16px);-webkit-backdrop-filter:blur(16px);border-radius:var(--radius);padding:26px 28px;margin-bottom:18px;box-shadow:var(--shadow);border:1px solid var(--glass-line)}
  .card-title{font-size:12px;font-weight:700;color:var(--mut-2);text-transform:uppercase;letter-spacing:.08em;margin:0 0 6px}
  .card h2{font-size:22px;font-weight:700;margin:0 0 14px;color:var(--txt);display:flex;align-items:center;gap:10px}
  .card h2 .ic{width:32px;height:32px;border-radius:10px;background:var(--teal-bg);color:var(--teal-d);display:inline-flex;align-items:center;justify-content:center;font-size:16px}

  /* ===== KPI 卡片(图片风格：图标徽章 + 标题 + 描述 + 大数字 + 脚注) ===== */
  .kpis{display:grid;grid-template-columns:repeat(auto-fit,minmax(190px,1fr));gap:14px;margin-bottom:16px}
  .kpi{background:var(--glass);backdrop-filter:blur(16px);-webkit-backdrop-filter:blur(16px);border-radius:var(--radius);padding:18px 20px;box-shadow:var(--shadow-sm);border:1px solid var(--glass-line);position:relative;display:flex;flex-direction:column;gap:8px;min-height:124px}
  .kpi-head{display:flex;align-items:flex-start;justify-content:space-between;gap:10px}
  .kpi-ic{width:40px;height:40px;border-radius:12px;display:inline-flex;align-items:center;justify-content:center;font-size:20px;color:#fff;background:var(--teal);box-shadow:0 4px 10px rgba(43,127,255,.3);flex-shrink:0}
  .kpi.sun .kpi-ic{background:var(--coral);box-shadow:0 4px 10px rgba(244,124,107,.32);color:#fff}
  .kpi.pink .kpi-ic{background:var(--good);box-shadow:0 4px 10px rgba(93,190,108,.3);color:#fff}
  .kpi.lav .kpi-ic{background:var(--lav);box-shadow:0 4px 10px rgba(205,180,240,.5);color:#6a4f9a}
  .kpi.sky .kpi-ic{background:var(--warn);box-shadow:0 4px 10px rgba(245,158,11,.3);color:#fff}
  .kpi.rose .kpi-ic{background:var(--rose);box-shadow:0 4px 10px rgba(245,163,163,.5);color:#a06464}
  .kpi-meta{text-align:right;flex:1;min-width:0}
  .kpi-tag{font-size:11px;font-weight:700;color:var(--mut-2);letter-spacing:.04em}
  .kpi-trend{font-size:11px;font-weight:700;color:var(--teal);margin-top:2px}
  .kpi-trend.down{color:var(--rose)}
  .kpi-title{font-size:14px;font-weight:700;color:var(--txt);margin:0;line-height:1.3}
  .kpi-sub{font-size:12px;color:var(--mut);margin-top:-2px;line-height:1.3}
  .kpi-value{font-size:36px;font-weight:800;letter-spacing:-.03em;color:var(--txt);line-height:1.1;font-variant-numeric:tabular-nums;margin-top:2px}
  .kpi-value .unit{font-size:14px;font-weight:700;color:var(--mut);margin-left:4px}
  .kpi-foot{font-size:12px;color:var(--mut);font-weight:500;margin-top:auto;padding-top:6px;border-top:1px dashed var(--line-soft);line-height:1.4}
  .kpi-foot b{color:var(--txt);font-weight:700}

  /* ===== 周图表风格"突出卡片"：天蓝/青绿实色底 ===== */
  .feature-card{background:linear-gradient(135deg,#a8e0ea,#7dd3e5);color:#1f2a3d;border-radius:var(--radius);padding:20px 22px;box-shadow:0 8px 24px rgba(125,211,229,.28);position:relative;overflow:hidden}
  .feature-card .fc-head{display:flex;align-items:center;justify-content:space-between;margin-bottom:14px}
  .feature-card .fc-title{font-size:13px;font-weight:700;color:#3a6a85;letter-spacing:.04em}
  .feature-card .fc-pills{display:flex;gap:6px}
  .feature-card .fc-pill{background:rgba(255,255,255,.6);color:#3a6a85;font-size:11px;font-weight:700;padding:3px 10px;border-radius:8px}
  .feature-card .fc-bars{display:flex;align-items:flex-end;gap:6px;height:80px;padding:8px 0}
  .feature-card .fc-bar{flex:1;background:#6a8aa8;border-radius:4px 4px 2px 2px;position:relative;min-height:6px}
  .feature-card .fc-bar.alt{background:rgba(255,255,255,.75)}
  .feature-card .fc-labels{display:flex;gap:6px;margin-top:6px;font-size:10.5px;color:#3a6a85;font-weight:600}
  .feature-card .fc-labels span{flex:1;text-align:center}
  .feature-card .fc-num{display:flex;align-items:baseline;gap:18px;margin-bottom:4px}
  .feature-card .fc-num .n{font-size:26px;font-weight:800;color:#2b3a4a;letter-spacing:-.02em;font-variant-numeric:tabular-nums}
  .feature-card .fc-num .n small{font-size:13px;color:#075985;font-weight:600;margin-left:3px}

  /* ===== 类别条形图(简化为青绿色阶) ===== */
  .cat-bars{background:var(--card);border-radius:var(--radius);padding:20px 22px;box-shadow:var(--shadow-sm);border:1px solid var(--line);margin-bottom:16px}
  .cat-bars h3{font-size:15px;font-weight:700;margin:0 0 14px;color:var(--txt);display:flex;align-items:center;gap:8px}
  .cat-bars h3 .ic{width:24px;height:24px;border-radius:8px;background:var(--teal-bg);color:var(--teal);display:inline-flex;align-items:center;justify-content:center;font-size:13px}
  .bar-row{display:grid;grid-template-columns:72px 1fr 48px;gap:12px;align-items:center;margin-bottom:8px;font-size:13px}
  .bar-row .name{color:var(--mut);font-weight:600;text-align:right;font-size:12.5px}
  .bar-row .num{color:var(--txt);font-weight:700;font-variant-numeric:tabular-nums;text-align:right;font-size:13px}
  .bar-track{height:8px;background:var(--line-soft);border-radius:6px;overflow:hidden}
  .bar-fill{height:100%;border-radius:6px;transition:width .6s cubic-bezier(.4,0,.2,1);background:var(--teal)}
  /* ===== 分类浓度总量（单报告） ===== */
  .conc-rows{margin-top:4px}
  .conc-row{display:grid;grid-template-columns:104px 1fr 96px 56px;gap:12px;align-items:center;margin-bottom:9px;font-size:13px}
  .conc-row .cb{display:flex;align-items:center;gap:7px;min-width:0}
  .conc-row .ct{height:11px;background:var(--line-soft);border-radius:6px;overflow:hidden}
  .conc-row .cf{height:100%;border-radius:6px;background:linear-gradient(90deg,var(--teal),var(--teal-d));transition:width .6s cubic-bezier(.4,0,.2,1)}
  .conc-row .cv{color:var(--txt);font-weight:800;font-variant-numeric:tabular-nums;text-align:right;font-size:13.5px}
  .conc-row .cv small{font-size:11px;color:var(--mut);font-weight:600;margin-left:2px}
  .conc-row .cp{color:var(--mut);font-weight:700;text-align:right;font-variant-numeric:tabular-nums;font-size:12.5px}
  .conc-foot{margin-top:12px;font-size:12.5px;color:var(--mut);display:flex;gap:16px;flex-wrap:wrap}
  .conc-foot b{color:var(--txt);font-weight:800}
  /* 类别卡片网格（图标 + 类别名 + 数量，按数量排序，可点击） */
  .cat-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(118px,1fr));gap:10px;margin-top:2px}
  .cat-chip{cursor:pointer;background:var(--teal-bg);border:1px solid var(--teal-line,var(--line));border-radius:12px;padding:12px 13px;display:flex;flex-direction:column;gap:9px;transition:transform .15s ease,box-shadow .15s ease,background .15s ease;position:relative;min-height:78px}
  .cat-chip:hover{transform:translateY(-2px);box-shadow:0 6px 16px rgba(125,211,229,.28);background:#e6f6fa}
  .cat-chip .cc-head{display:flex;align-items:center;gap:7px;min-width:0}
  .cat-chip .cc-dot{width:11px;height:11px;border-radius:50%;flex:0 0 auto;box-shadow:0 1px 3px rgba(0,0,0,.18)}
  .cat-chip .cc-dot.b-醛类{background:var(--c-醛类)} .cat-chip .cc-dot.b-酮类{background:var(--c-酮类)}
  .cat-chip .cc-dot.b-酯类{background:var(--c-酯类)} .cat-chip .cc-dot.b-酸类{background:var(--c-酸类)}
  .cat-chip .cc-dot.b-醇类{background:var(--c-醇类)} .cat-chip .cc-dot.b-萜烯类与含氧萜类{background:var(--c-萜烯类与含氧萜类)}
  .cat-chip .cc-dot.b-含硫化合物{background:var(--c-含硫化合物)} .cat-chip .cc-dot.b-吡嗪类{background:var(--c-吡嗪类)}
  .cat-chip .cc-dot.b-内酯类{background:var(--c-内酯类)} .cat-chip .cc-dot.b-其他{background:var(--c-其他)}
  .cat-chip .cc-dot.b-含氮杂环其他含氮化合物{background:var(--c-含氮杂环其他含氮化合物)}
  .cat-chip .cc-dot.b-呋喃呋喃酮类{background:var(--c-呋喃呋喃酮类)}
  .cat-chip .cc-dot.b-噻唑类噻唑啉类{background:var(--c-噻唑类噻唑啉类)}
  .cat-chip .cc-name{font-size:13.5px;font-weight:700;color:var(--txt);letter-spacing:-.01em;line-height:1.25;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
  .cat-chip .cc-num{margin-top:auto;font-size:24px;font-weight:800;color:var(--teal-d);letter-spacing:-.03em;font-variant-numeric:tabular-nums;line-height:1;display:flex;align-items:baseline;gap:3px}
  .cat-chip .cc-num small{font-size:12px;font-weight:700;color:var(--mut)}
  .cat-row.is-active{background:#e6f6fa;box-shadow:0 6px 16px rgba(125,211,229,.34);border-color:var(--teal)}
  .cat-detail{margin-top:14px}
  .cat-detail:empty{margin-top:0}
  .cat-detail .card{margin:0}
  /* 多报告对比工具条 */
  .multi-toolbar{display:flex;align-items:center;gap:14px;flex-wrap:wrap;margin:4px 0 14px}
  .multi-toolbar label{font-size:13px;color:var(--mut);font-weight:600}
  .multi-toolbar select{padding:7px 10px;border-radius:8px;border:1px solid var(--line);background:var(--card);color:var(--txt);font-size:13px;margin-left:6px}

  /* ===== Top 排行卡片(图片风格：圆序号+名称+大数字) ===== */
  .top-list{background:var(--card);border-radius:var(--radius);padding:20px 22px;box-shadow:var(--shadow-sm);border:1px solid var(--line);margin-bottom:16px}
  .top-list h3{font-size:15px;font-weight:700;margin:0 0 14px;color:var(--txt);display:flex;align-items:center;gap:8px}
  .top-list h3 .ic{width:24px;height:24px;border-radius:8px;background:var(--sun-bg);color:var(--sun);display:inline-flex;align-items:center;justify-content:center;font-size:13px}
  .top-item{display:flex;align-items:center;gap:12px;padding:10px 0;border-bottom:1px solid var(--line-soft)}
  .top-item:last-child{border-bottom:0}
  .top-rank{width:30px;height:30px;border-radius:50%;background:var(--teal-bg);color:var(--teal);font-weight:800;font-size:13px;display:flex;align-items:center;justify-content:center;flex-shrink:0;font-variant-numeric:tabular-nums}
  .top-item:nth-child(2) .top-rank{background:var(--sun-bg);color:var(--sun)}
  .top-item:nth-child(3) .top-rank{background:var(--sky-bg);color:var(--sky-d)}
  .top-name{flex:1;min-width:0}
  .top-name .en{font-weight:700;color:var(--txt);font-size:14px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
  .top-name .cn{font-size:12.5px;color:var(--mut);font-weight:600;margin-top:1px}
  .top-val{font-size:20px;font-weight:800;color:var(--txt);font-variant-numeric:tabular-nums;flex-shrink:0;letter-spacing:-.02em}
  .top-val small{font-size:11px;color:var(--mut);font-weight:500;margin-left:3px}

  /* ===== 表单 ===== */
  textarea,input[type=text],input[type=file]{width:100%;background:var(--card-soft);border:1.5px solid var(--line);color:var(--txt);border-radius:var(--radius-sm);padding:12px 14px;font-size:14px;font-family:inherit;transition:all .2s}
  textarea:focus,input[type=text]:focus,input[type=file]:focus{outline:none;border-color:var(--teal);box-shadow:0 0 0 4px var(--teal-bg);background:#fff}
  textarea{min-height:120px;resize:vertical}
  select{font-family:inherit;font-size:14px;padding:10px 14px;border-radius:10px;border:1.5px solid var(--line);background:var(--card-soft);color:var(--txt);cursor:pointer;font-weight:600}
  .row{display:flex;gap:12px;flex-wrap:wrap;align-items:center;margin-top:14px}
  .row>*{flex:1;min-width:0}

  /* 参数网格：标签在上、输入在下，整齐对齐（消除一上一下） */
  .fields{display:grid;grid-template-columns:repeat(auto-fit,minmax(120px,1fr));gap:12px;margin-top:14px}
  .field{display:flex;flex-direction:column;gap:6px}
  .field>span{font-size:12px;font-weight:600;color:#5aa9da;letter-spacing:.01em}
  .field input{width:100%}
  .conc-note{margin-top:12px;font-size:12px;color:var(--mut);line-height:1.5;padding:10px 12px;background:var(--card-soft);border-radius:var(--radius-sm);border:1px solid var(--line)}

  /* 上传区：清晰卡片化，杜绝错落（粉蓝系精确配色） */
  .upload-zone{display:flex;align-items:center;gap:14px;background:var(--card-soft);border:1.5px dashed #c3dde6;border-radius:var(--radius-sm);padding:14px 16px;transition:all .2s;margin-top:12px}
  .upload-zone:hover{border-color:var(--teal);background:var(--teal-bg)}
  .upload-zone .up-ic{width:42px;height:42px;border-radius:12px;background:var(--coral);color:#fff;display:inline-flex;align-items:center;justify-content:center;font-size:20px;flex-shrink:0;box-shadow:0 4px 10px rgba(244,124,107,.35)}
  .upload-zone .up-body{flex:1;min-width:0;display:flex;flex-direction:column;gap:2px}
  .upload-zone .up-title{font-size:13.5px;font-weight:700;color:var(--txt)}
  .upload-zone .up-sub{font-size:12px;color:#6b7889}
  .upload-zone input[type=file]{display:none}
  .upload-zone .up-btn{background:var(--teal);color:#1f2a3d;border:0;border-radius:10px;padding:9px 16px;font-size:13px;font-weight:700;cursor:pointer;white-space:nowrap;transition:all .2s;box-shadow:0 3px 10px rgba(125,211,229,.35)}
  .upload-zone .up-btn:hover{background:var(--coral);color:#fff}

  /* ===== 按钮(信任蓝主色 + 珊瑚橙 hover) ===== */
  button{background:var(--teal);color:#1f2a3d;border:0;border-radius:var(--radius-sm);padding:11px 22px;font-size:14px;font-weight:700;cursor:pointer;transition:all .2s;box-shadow:0 4px 12px rgba(125,211,229,.35);letter-spacing:.02em}
  button:hover{background:var(--coral);color:#fff;transform:translateY(-1px);box-shadow:0 6px 18px rgba(244,124,107,.35)}
  button:active{transform:translateY(0)}
  button.ghost{background:#fff;border:1.5px solid var(--line);color:var(--teal-d);box-shadow:none}
  button.ghost:hover{background:var(--teal-bg);border-color:var(--teal);color:var(--teal-d)}

  .hint{color:var(--mut);font-size:13px;margin-top:8px;line-height:1.5}
  .hint b{color:var(--txt)}
  .panel{display:none}
  .panel.active{display:block;animation:fade .25s ease}
  @keyframes fade{from{opacity:0;transform:translateY(6px)}to{opacity:1;transform:none}}

  /* ===== 全局数据表设计规范（记忆格式 · 所有数据表统一套用） ===== */
  /* 容器：圆角 + 细边框 + 轻投影 + 可滚动 + 滚动条美化 */
  .wrap{overflow:auto;max-height:66vh;border-radius:var(--radius);border:1px solid var(--line);background:#fff;box-shadow:var(--shadow-sm)}
  .wrap::-webkit-scrollbar{height:10px;width:10px}
  .wrap::-webkit-scrollbar-thumb{background:#d8dee6;border-radius:8px;border:2px solid #fff}
  .wrap::-webkit-scrollbar-thumb:hover{background:#c2ccd6}
  /* 基础表格：行 / 列 / 文字分布适中，清晰易读 */
  table{border-collapse:separate;border-spacing:0;width:auto;min-width:100%;font-size:12.5px;line-height:1.55;margin:0;color:var(--txt)}
  th,td{padding:10px 14px;text-align:left;vertical-align:top;border-bottom:1px solid var(--line-soft);word-break:break-word;overflow-wrap:anywhere}
  /* 表头：浅底 + 加粗小字 + 吸顶 + 加重底线 + 圆角裁切 */
  th{background:var(--card-soft);color:var(--mut);font-weight:700;font-size:11px;position:sticky;top:0;z-index:2;border-bottom:1.5px solid var(--line);white-space:nowrap;letter-spacing:.03em}
  .wrap thead th:first-child{border-top-left-radius:var(--radius)}
  .wrap thead th:last-child{border-top-right-radius:var(--radius)}
  /* 斑马纹 + hover 高亮，末行去线 */
  tbody tr:nth-child(even) td{background:#fcfdfe}
  tbody tr:hover td{background:#eef5fb}
  tr:last-child td{border-bottom:0}
  /* 数值列：右对齐 + 等宽数字，整数位对齐清晰 */
  td.num,th.num{text-align:right;font-variant-numeric:tabular-nums;white-space:nowrap}
  td .cn{font-weight:700;color:var(--txt);font-size:13.5px}
  /* 描述 / 来源 列：行高舒展、字号适中、长文本换行不挤压（title 显示全文） */
  td.desc,td.gc-desc{text-align:left;white-space:normal;color:var(--txt);font-size:12px;line-height:1.55;max-width:240px}
  td.src,td.gc-src{text-align:left;white-space:normal;color:var(--mut);font-size:12px;line-height:1.6;max-width:180px}
  td.src b,td.gc-src b{color:var(--txt);font-weight:700}
  /* 报告解析主表：固定列宽，压缩中英文/物质列，保证后续数据列·来源·风味描述清晰 */
  #report-table{table-layout:fixed;width:100%}
  #report-table th,#report-table td{font-variant-numeric:tabular-nums}
  #report-table th:nth-child(1),#report-table td:nth-child(1){width:34px;text-align:center;white-space:nowrap;color:var(--mut-2);font-variant-numeric:tabular-nums}
  #report-table th:nth-child(2),#report-table td:nth-child(2){width:13%;min-width:104px;text-align:left;white-space:normal;word-break:break-word;overflow-wrap:anywhere}
  #report-table th:nth-child(3),#report-table td:nth-child(3){width:8%;min-width:68px;text-align:left;white-space:normal;word-break:break-word}
  #report-table th:nth-child(4),#report-table td:nth-child(4){text-align:left;white-space:nowrap;min-width:80px}
  #report-table th:nth-child(5),#report-table td:nth-child(5){width:8%;min-width:58px;text-align:left;white-space:normal}
  #report-table th:nth-child(6),#report-table th:nth-child(7),#report-table th:nth-child(10),#report-table th:nth-child(11),#report-table th:nth-child(12),
  #report-table td:nth-child(6),#report-table td:nth-child(7),#report-table td:nth-child(10),#report-table td:nth-child(11),#report-table td:nth-child(12){text-align:right;white-space:nowrap}
  #report-table th:nth-child(8),#report-table th:nth-child(9),
  #report-table td:nth-child(8),#report-table td:nth-child(9){text-align:right;white-space:normal}
  #report-table th:nth-child(12),#report-table td:nth-child(12){text-align:center;white-space:nowrap;min-width:88px}
  #report-table th:nth-child(13),#report-table td:nth-child(13){width:19%;min-width:196px;text-align:left;white-space:normal}
  #report-table th:nth-child(14),#report-table td:nth-child(14){width:13%;min-width:150px;text-align:left;white-space:normal}
  /* 数据库浏览表：固定列宽布局，长文本/长类别优雅换行不挤压、不叠放 */
  #browse-wrap table{table-layout:fixed;width:100%}
  #browse-wrap th,#browse-wrap td{word-break:break-word;overflow-wrap:anywhere}
  #browse-wrap .c-idx{width:42px;text-align:center;white-space:nowrap;color:var(--mut-2);font-variant-numeric:tabular-nums}
  #browse-wrap .c-en{width:16%}
  #browse-wrap .c-cn{width:11%}
  #browse-wrap .c-cas{width:10%}
  #browse-wrap .c-cat{width:12%}
  #browse-wrap .c-thr{width:9%}
  #browse-wrap .c-med{width:7%}
  #browse-wrap .c-match{width:6%}
  #browse-wrap .c-odor{width:19%}
  #browse-wrap .c-src{width:11%}
  /* 气味/来源描述列：行高舒展、字号适中，多行不拥挤 */
  #browse-wrap td.c-odor,#browse-wrap td.c-src{color:var(--mut);line-height:1.6;font-size:12px;vertical-align:top}
  #browse-wrap td.c-odor{color:var(--txt)}
  /* 类别徽章在窄列内允许换行，避免长类别名溢出叠放 */
  #browse-wrap td.c-cat{padding:8px 8px}
  #browse-wrap td.c-cat .badge{white-space:normal;line-height:1.3;font-size:10.5px;text-align:center;display:inline-block;max-width:100%}
  /* 数据库中文名加粗、突出 */
  #browse-wrap td.c-cn{font-weight:800;color:var(--txt);font-size:13px}
  /* 斑马纹：隔行浅底，长表更易读（hover 仍可见） */
  #browse-wrap tbody tr:nth-child(even) td{background:#fcfdfe}
  #browse-wrap tbody tr:hover td{background:#eef5fb}
  /* 多报告对比矩阵：首列固定宽，样品列等宽 */
  #multi-matrix table{table-layout:fixed;width:100%}
  #multi-matrix th,#multi-matrix td{word-break:break-word;vertical-align:top}
  #multi-matrix th:first-child,#multi-matrix td:first-child{width:16%}
  #multi-matrix th:nth-child(2),#multi-matrix td:nth-child(2){width:11%}
  #multi-matrix th:nth-child(3),#multi-matrix td:nth-child(3){width:8%}
  #multi-matrix td.m-none,#multi-matrix th:last-child{text-align:center;color:var(--mut-2)}
  #multi-matrix td small{color:var(--mut);font-weight:500}
  /* 相对于对照列：上调(≥2×)红、下调(≤0.5×)蓝，与单 PDF 一致 */
  #multi-matrix .fc-up{color:#e0533a;font-weight:700}
  #multi-matrix .fc-down{color:#2a7de1;font-weight:700}
  /* OAV 气味活性标志 */
  .oavflag{font-weight:700;font-size:12.5px;white-space:nowrap}
  .oavflag.f-key{color:var(--good);background:var(--good-bg);padding:2px 8px;border-radius:8px}
  .oavflag.f-pot{color:var(--coral);background:var(--coral-bg);padding:2px 8px;border-radius:8px}
  .oavflag.f-na{color:var(--mut-2)}
  /* 矩阵数值加深（清晰可读） */
  .num-dark{color:var(--txt);font-weight:700;font-variant-numeric:tabular-nums;font-size:13px;line-height:1.5}
  #multi-matrix td.oavflag.num-dark small{color:var(--mut);font-weight:600}
  /* 各样品分类浓度总量对比（多报告） */
  .multi-cat{width:100%;border-collapse:collapse}
  .multi-cat th,.multi-cat td{padding:9px 12px;text-align:right;border-bottom:1px solid var(--line-soft);font-variant-numeric:tabular-nums;white-space:nowrap}
  .multi-cat th:first-child,.multi-cat td:first-child{text-align:left}
  .multi-cat thead th{background:var(--card-soft);color:var(--mut);font-weight:700;font-size:11px;position:sticky;top:0;z-index:2;border-bottom:1.5px solid var(--line)}
  .multi-cat .cc-cell{position:relative}
  .multi-cat .cc-bar{position:absolute;left:12px;top:50%;transform:translateY(-50%);height:15px;border-radius:4px;background:linear-gradient(90deg,var(--teal),var(--teal-d));opacity:.22;min-width:2px;max-width:calc(100% - 24px)}
  .multi-cat .cc-v{position:relative;font-weight:700}
  .multi-cat .cc-sum{font-weight:800;color:var(--teal-d)}
  .multi-cat .cc-total td{border-top:2px solid var(--line);background:var(--card-soft);font-weight:800}

  /* ===== Nature 风格可视化（热力图 / PCA）+ 物质管理 ===== */
  .viz-wrap{margin-top:18px;border-top:1px dashed var(--line);padding-top:14px}
  .viz-head{display:flex;align-items:center;gap:10px;flex-wrap:wrap;margin:4px 0 10px}
  .viz-head h3{margin:0;font-size:15px;display:flex;align-items:center;gap:8px}
  .viz-head h3 .ic{width:24px;height:24px;border-radius:8px;background:var(--teal-bg);color:var(--teal);display:inline-flex;align-items:center;justify-content:center;font-size:13px}
  .viz-toolbar{display:flex;flex-wrap:wrap;gap:8px;align-items:center;margin-bottom:10px}
  .viz-toolbar label{font-size:12.5px;color:var(--mut);font-weight:600}
  .viz-toolbar select{padding:6px 9px;border-radius:8px;border:1px solid var(--line);background:var(--card);color:var(--txt);font-size:13px}
  .viz-chips{display:flex;flex-wrap:wrap;gap:6px;margin:2px 0 8px}
  .viz-chip{display:inline-flex;align-items:center;gap:5px;padding:3px 10px;border-radius:18px;border:1px solid var(--line);background:var(--card);font-size:12px;font-weight:600;cursor:pointer;user-select:none;transition:all .15s}
  .viz-chip .dot{width:9px;height:9px;border-radius:50%}
  .viz-chip.off{opacity:.4;background:#f1f3f6}
  .viz-chip:hover{border-color:var(--teal)}
  .viz-chip em.cn{font-style:normal;font-size:10.5px;color:var(--mut);background:var(--card-soft);border-radius:8px;padding:0 6px;margin-left:2px;font-weight:600}
  .viz-chip.off em.cn{color:var(--mut-2)}
  .viz-search{flex:1;min-width:160px;padding:7px 10px;border-radius:9px;border:1px solid var(--line);background:#fff;color:var(--txt);font-size:13px}
  .viz-main{display:grid;grid-template-columns:minmax(320px,1.4fr) minmax(300px,1fr);gap:14px;align-items:stretch}
  @media(max-width:900px){ .viz-main{grid-template-columns:1fr} .viz-mgr{max-height:520px} }
  .viz-fig{background:#fff;border:1px solid var(--line);border-radius:var(--radius);padding:12px 12px 8px;box-shadow:var(--shadow-sm)}
  .viz-fig h4{margin:0 0 6px;font-size:13px;color:var(--txt);font-weight:700;display:flex;justify-content:space-between;align-items:baseline}
  .viz-fig h4 small{font-weight:500;color:var(--mut);font-size:11px}
  .viz-fig svg{width:100%;height:auto;display:block;font-family:var(--sans)}
  .viz-mgr{background:#fff;border:1px solid var(--line);border-radius:var(--radius);padding:10px 12px;box-shadow:var(--shadow-sm);display:flex;flex-direction:column}
  .viz-mgr .mgr-top{display:flex;gap:6px;flex-wrap:wrap;margin-bottom:8px}
  .viz-mgr .mgr-list{overflow:auto;flex:1;border-top:1px solid var(--line-soft)}
  .viz-mrow{display:flex;align-items:center;gap:8px;padding:5px 6px;border-bottom:1px solid var(--line-soft);font-size:12.5px}
  .viz-mrow:hover{background:#f6fafc}
  .viz-mrow input{width:15px;height:15px;accent-color:var(--teal);flex:0 0 auto}
  .viz-mrow .mn{flex:1;min-width:0;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
  .viz-mrow .mn b{color:var(--txt)}
  .viz-mrow .mc{font-size:10.5px;color:var(--mut);flex:0 0 auto;padding:1px 7px;border-radius:7px;background:var(--card-soft)}
  .viz-mrow .mv{font-size:11px;color:var(--mut-2);font-variant-numeric:tabular-nums;flex:0 0 auto;min-width:54px;text-align:right}
  .viz-mrow.del{opacity:.45}
  .viz-mrow.del .mn{text-decoration:line-through}
  .viz-hint{font-size:11.5px;color:var(--mut-2);margin-top:6px;line-height:1.5}
  .viz-legend{display:flex;flex-wrap:wrap;gap:10px;margin-top:6px;font-size:11px;color:var(--mut)}
  .viz-legend span{display:inline-flex;align-items:center;gap:4px}
  .viz-legend i{width:11px;height:11px;border-radius:3px;display:inline-block}
  .viz-sep{flex:0 0 auto;width:1px;height:22px;background:var(--line);margin:0 2px}
  .viz-toolbar button.ghost[id^="viz-png-"],.viz-toolbar button.ghost[id^="viz-svg-"]{border-color:var(--teal);color:var(--teal-d)}
  .viz-axis{stroke:#9aa6b2;stroke-width:1}
  .viz-axis-t{fill:#5b6776;font-size:11px}
  .viz-rlabel{fill:#2a3340;font-size:11px}
  .viz-clabel{fill:#5b6776;font-size:11px;font-weight:600}
  .viz-cbar-t{fill:#5b6776;font-size:10.5px}

  /* ===== 处理组对比（单 PDF 多处理组自动识别） ===== */
  .gc-banner{background:linear-gradient(90deg,var(--card-soft),var(--card));border:1px solid var(--line);border-left:4px solid var(--teal);border-radius:12px;padding:12px 16px;margin:6px 0 12px}
  .gc-banner b{color:var(--teal-d)}
  .gc-chips{margin-top:6px;display:flex;flex-wrap:wrap;gap:6px}
  .gc-chip{display:inline-flex;align-items:center;gap:6px;background:var(--card);border:1px solid var(--line);border-radius:20px;padding:3px 11px;font-size:12.5px;font-weight:700}
  .gc-chip .dot{width:9px;height:9px;border-radius:50%}
  .gc-summary{display:flex;flex-wrap:wrap;gap:10px;margin:10px 0}
  .gc-stat{flex:1;min-width:150px;background:var(--card);border:1px solid var(--line);border-radius:10px;padding:10px 13px}
  .gc-stat h4{margin:0 0 4px;font-size:12px;color:var(--mut);font-weight:700}
  .gc-stat .v{font-size:20px;font-weight:800;color:var(--teal-d)}
  .gc-stat .sub{font-size:11.5px;color:var(--mut-2)}
  /* 处理组对比矩阵：固定列宽，压缩「物质/中文名/类别」列，把宽度让给数据列与风味描述 */
  #gc-matrix{table-layout:fixed;width:100%}
  #gc-matrix th,#gc-matrix td{text-align:right;font-variant-numeric:tabular-nums;white-space:nowrap}
  /* 首列删除按钮：固定窄列、居中 */
  #gc-matrix th.gc-del-col,#gc-matrix td.gc-del-col{width:36px;text-align:center;white-space:normal}
  /* 物质(EN) / 中文名 / 类别：压缩并允许换行，避免长英文名撑爆列宽 */
  #gc-matrix th:nth-child(2),#gc-matrix td:nth-child(2){text-align:left;width:14%;min-width:120px;white-space:normal;word-break:break-word;overflow-wrap:anywhere}
  #gc-matrix th:nth-child(3),#gc-matrix td:nth-child(3){text-align:left;width:9%;min-width:76px;white-space:normal;word-break:break-word}
  #gc-matrix th:nth-child(4),#gc-matrix td:nth-child(4){text-align:left;width:8%;min-width:62px;white-space:normal}
  /* 各处理组数据列：等宽，内容堆叠不挤压 */
  #gc-matrix .grp-col{min-width:84px;background:var(--card-soft)}
  /* 相对对照列：左对齐、长文本换行 */
  #gc-matrix th:nth-last-child(2),#gc-matrix td:nth-last-child(2){width:11%;min-width:104px;text-align:left;white-space:normal}
  /* 风味描述列（末列）：占充足宽度，长文本换行不挤压 */
  #gc-matrix th:last-child,#gc-matrix td:last-child{text-align:left;white-space:normal;width:22%;min-width:220px}
  #gc-matrix td.gc-desc{color:var(--txt);font-size:12px;line-height:1.55;max-width:300px}
  #gc-matrix .grp-col{background:var(--card-soft)}
  #gc-matrix .cell-rr{font-weight:700;font-size:13px}
  #gc-matrix .cell-conc{font-size:11.5px;color:var(--mut)}
  #gc-matrix .cell-oav{font-size:11.5px}
  #gc-matrix td.hl-max{background:rgba(36,160,140,.14);box-shadow:inset 0 0 0 1px rgba(36,160,140,.35)}
  #gc-matrix .gc-up{color:#e0533a;font-weight:700}
  #gc-matrix .gc-down{color:#2a7de1;font-weight:700}
  #gc-matrix .gc-na{color:var(--mut-2)}
  #gc-matrix .fc-up{color:#e0533a;font-weight:700}
  #gc-matrix .fc-down{color:#2a7de1;font-weight:700}
  .gc-tools{display:flex;gap:8px;flex-wrap:wrap;margin:8px 0}
  .gc-mgrbar{display:flex;align-items:center;gap:8px;flex-wrap:wrap;background:var(--card);border:1px solid var(--line);border-radius:10px;padding:8px 12px;margin:8px 0}
  .gc-mgrbar>b{font-size:13px;color:var(--teal-d)}
  .gc-mgrbar .hint{font-size:11.5px}
  #gc-matrix th.gc-del-col,#gc-matrix td.gc-del-col{width:34px;text-align:center;padding:4px}
  #gc-matrix td.gc-del-col .gc-del-btn{width:22px;height:22px;border-radius:6px;border:1px solid var(--line);background:#fff;color:var(--rose,#d9534f);font-size:13px;font-weight:800;line-height:1;cursor:pointer;transition:all .15s}
  #gc-matrix td.gc-del-col .gc-del-btn:hover{background:var(--rose,#d9534f);color:#fff;border-color:var(--rose,#d9534f)}
  /* 多报告对比矩阵：每行删除（小圆点 ✕），与物质管理删除状态共享 */
  #multi-matrix th.multi-del-col,#multi-matrix td.multi-del-col{width:30px;text-align:center;padding:4px;white-space:normal}
  #multi-matrix td.multi-del-col .multi-del-btn{width:20px;height:20px;border-radius:50%;border:1px solid var(--line);background:#fff;color:var(--mut-2);font-size:12px;font-weight:800;line-height:1;cursor:pointer;transition:all .15s;display:inline-flex;align-items:center;justify-content:center;padding:0}
  #multi-matrix td.multi-del-col .multi-del-btn:hover{background:var(--rose,#d9534f);color:#fff;border-color:var(--rose,#d9534f)}
  #gc-stat-single{font-weight:700;color:var(--teal-d);font-size:12.5px}
  /* 主表类别筛选 */
  .gc-catbar{display:flex;flex-wrap:wrap;gap:6px;margin:2px 0 4px}
  .gc-cat-chip{padding:4px 12px;border-radius:18px;border:1px solid var(--line);background:var(--card);font-size:12px;font-weight:600;color:var(--teal-d);cursor:pointer;user-select:none;transition:all .15s}
  .gc-cat-chip.off{opacity:.45;background:#f1f3f6;color:var(--mut-2)}
  .gc-cat-chip:hover{border-color:var(--teal)}
  .gc-cat-chip:not(.off){background:var(--teal-bg);border-color:var(--teal)}
  /* 热力图物质管理：按类别分组 */
  .mgr-cat{margin-bottom:6px}
  .mgr-cat-h{position:sticky;top:0;display:flex;align-items:center;gap:6px;padding:5px 6px;font-size:12px;font-weight:800;color:var(--teal-d);background:var(--card);border-radius:6px;cursor:pointer;user-select:none}
  .mgr-cat-h:hover{background:#f1f6f8}
  .mgr-cat-h .mgr-chev{display:inline-flex;width:14px;justify-content:center;color:var(--mut-2);font-size:11px;transition:transform .15s}
  .mgr-cat.collapsed .mgr-cat-h .mgr-chev{transform:rotate(-90deg)}
  .mgr-cat-h .dot{width:9px;height:9px;border-radius:50%}
  .mgr-cat-h em.cn{font-style:normal;font-size:10.5px;color:var(--mut);margin-left:auto;font-weight:600}
  .mgr-cat.collapsed .mgr-cat-body{display:none}

  /* 类别徽章：浅底深字 */
  .badge{display:inline-block;padding:3px 10px;border-radius:8px;font-size:11.5px;font-weight:700;white-space:nowrap;letter-spacing:.02em;background:var(--cb-其他);color:var(--c-其他)}
  .b-醛类{background:var(--cb-醛类);color:var(--c-醛类)}
  .b-酮类{background:var(--cb-酮类);color:var(--c-酮类)}
  .b-酯类{background:var(--cb-酯类);color:var(--c-酯类)}
  .b-酸类{background:var(--cb-酸类);color:var(--c-酸类)}
  .b-醇类{background:var(--cb-醇类);color:var(--c-醇类)}
  .b-萜烯类与含氧萜类{background:var(--cb-萜烯类与含氧萜类);color:var(--c-萜烯类与含氧萜类)}
  .b-含硫化合物{background:var(--cb-含硫化合物);color:var(--c-含硫化合物)}
  .b-吡嗪类{background:var(--cb-吡嗪类);color:var(--c-吡嗪类)}
  .b-内酯类{background:var(--cb-内酯类);color:var(--c-内酯类)}
  .b-含氮杂环其他含氮化合物{background:var(--cb-含氮杂环其他含氮化合物);color:var(--c-含氮杂环其他含氮化合物)}
  .b-呋喃呋喃酮类{background:var(--cb-呋喃呋喃酮类);color:var(--c-呋喃呋喃酮类)}
  .b-噻唑类噻唑啉类{background:var(--cb-噻唑类噻唑啉类);color:var(--c-噻唑类噻唑啉类)}
  .b-其他{background:var(--cb-其他);color:var(--c-其他)}

  .stat{display:flex;gap:18px;flex-wrap:wrap;margin-top:14px;font-size:13px;color:var(--mut);align-items:center}
  .stat b{color:var(--txt);font-weight:700;font-size:15px}
  .pill{background:var(--teal-bg);color:var(--teal-d);padding:4px 12px;border-radius:8px;font-size:12.5px;font-weight:700}
  .ok{color:var(--teal);font-weight:700}.warn{color:var(--sun);font-weight:700}.bad{color:var(--rose);font-weight:700}

  /* KPI 网格在窄屏自适应 */
  @media (max-width:760px){
    h1{font-size:24px}
    .kpi-value{font-size:26px}
    .kpis{grid-template-columns:repeat(2,1fr)}
    .card{padding:18px 20px}
    table th,table td{font-size:12px;padding:8px 9px}
    .bar-row{grid-template-columns:54px 1fr 48px}
  }
  @media (max-width:480px){
    .kpis{grid-template-columns:1fr 1fr}
    .bar-row{grid-template-columns:48px 1fr 40px;font-size:12px}
  }
</style>
</head>
<body class="notranslate" translate="no">
<h1>风味物质自动检索 · 分类 · 阈值</h1>
<div class="sub">内置 <b id="dbsize"></b> 种风味化合物数据库 · 支持英文/中文检索、模糊匹配、未知物按命名结构自动分类 · 阈值/浓度单位 μg/L（水相近似）</div>

<div class="tabs">
  <div class="tab active" data-t="single">单物质检索</div>
  <div class="tab" data-t="batch">批量检索</div>
  <div class="tab" data-t="report">GC-MS 报告解析</div>
  <div class="tab" data-t="multi">多报告对比</div>
  <div class="tab" data-t="browse">浏览数据库</div>
</div>

<!-- 单物质 -->
<div class="panel active" id="p-single">
  <div class="card">
    <input type="text" id="q" placeholder="输入物质英文名或中文名，如 Butanoic acid / 丁酸 / 2-Heptanone">
    <div class="row">
      <button onclick="doSingle()">检索</button>
      <span class="hint">回车也可检索</span>
    </div>
    <div id="single-out"></div>
  </div>
</div>

<!-- 批量 -->
<div class="panel" id="p-batch">
  <div class="card">
    <textarea id="batch" placeholder="每行一个物质名，例如：&#10;Butanoic acid&#10;2-Heptanone&#10;gamma-Nonalactone&#10;丁酸&#10;未知物 ABC"></textarea>
    <div class="row">
      <button onclick="doBatch()">批量检索</button>
      <button class="ghost" onclick="exportCSV('batch')">导出 CSV</button>
    </div>
    <label class="upload-zone">
      <span class="up-ic">⬆</span>
      <span class="up-body">
        <span class="up-title">上传物质列表文件</span>
        <span class="up-sub">支持 .txt / .csv / .xlsx / .pdf / .docx</span>
      </span>
      <span class="up-btn">选择文件</span>
      <input type="file" id="batchfile" accept=".txt,.csv,.xlsx,.xls,.pdf,.docx" onchange="loadBatchFile()">
    </label>
    <div id="batch-out"></div>
  </div>
</div>

<!-- 报告 -->
<div class="panel" id="p-report">
  <div class="card">
    <p class="hint">上传 GC-MS 定量报告（.xlsx）或文献/报告文档（.pdf / .docx / .md / .txt）：自动识别其中的风味物质，提取 化合物 / 类别 / 气味描述 / 来源，并按类别聚合、组内按气味活性排序。</p>
    <label class="upload-zone">
      <span class="up-ic">⬆</span>
      <span class="up-body">
        <span class="up-title">上传 GC-MS 报告 / 文献文档</span>
        <span class="up-sub">支持 .xlsx / .pdf / .docx / .md / .txt</span>
      </span>
      <span class="up-btn">选择文件</span>
      <input type="file" id="reportfile" accept=".xlsx,.xls,.pdf,.docx,.md,.markdown,.txt,.csv">
    </label>
    <div class="row" style="margin-top:12px">
      <button onclick="doReport()">解析并检索</button>
      <button class="ghost" onclick="exportCSV('report')">导出 CSV</button>
      <span class="hint" id="report-fileinfo" style="flex:2;margin-top:0"></span>
    </div>
    <div class="fields" style="margin-top:18px;padding-top:16px;border-top:1px solid var(--line)">
      <div class="field"><span>内标物质</span><input type="text" id="r-ist" value="2-辛醇" oninput="document.getElementById('r-ist').dataset.touched='1';liveRecalcReport()"></div>
      <div class="field"><span>内标浓度</span><input type="text" id="r-cis" value="10" oninput="liveRecalcReport()"></div>
      <div class="field"><span>浓度单位</span><input type="text" id="r-cisu" value="mg/L" oninput="liveRecalcReport()"></div>
      <div class="field"><span>内标加量</span><input type="text" id="r-vis" value="50" oninput="liveRecalcReport()"></div>
      <div class="field"><span>体积单位</span><input type="text" id="r-visu" value="μL" oninput="liveRecalcReport()"></div>
      <div class="field"><span>样品量</span><input type="text" id="r-ms" value="2" oninput="liveRecalcReport()"></div>
      <div class="field"><span>取样量单位</span><input type="text" id="r-msu" value="g" oninput="liveRecalcReport()"></div>
      <div class="field" style="justify-content:flex-end"><span>&nbsp;</span><button class="ghost" onclick="recalcReportConc()">重算浓度</button></div>
    </div>
    <div class="conc-note">内标参数（内标物质 / 浓度 / 加量 / 样品量）均可手动修改；任一参数变动即实时重算浓度与 OAV/ROAV，无需点「重算浓度」。<b>内标物质</b>会优先用报告自动识别值预填。</div>
    <div class="conc-note">按报告表头语义自动识别各列并映射到内标法公式字母：<b>A</b>=化合物峰面积(响应/峰面积)、<b>A₁</b>=内标峰面积(ISTD/内标)、<b>rr</b>=响应比=A÷A₁(无量纲)、<b>c</b>=浓度(报告自带)、<b>RT</b>=保留时间；<b>c₁</b>(内标浓度)/<b>V₁</b>(内标加量)/<b>m</b>(样品量) 为下方参数。浓度直接使用响应比计算：C = c₁·V₁·rr÷m，因 rr 已含内标峰面积故不再除以 A₁，结果换算到 μg/L；OAV=浓度÷阈值(>1 具气味活性)，ROAV=OAV÷最大OAV×100(≥10 关键致香，≥1 潜在贡献)。</div>
    <div id="report-out"></div>
  </div>
</div>

<!-- 多报告对比 -->
<div class="panel" id="p-multi">
  <div class="card">
    <p class="hint">多报告批量分析对比：逐个上传 GC-MS 报告（.xlsx / .pdf / .docx / .txt），每个文件上传后填入<b>样品名称</b>；全部添加后点「开始对比分析」，系统将统一按内标法计算各样品浓度/OAV/ROAV，并横向对比共有/特有物质与关键致香物。</p>
    <div class="row">
      <label class="upload-zone">
        <span class="up-ic">⬆</span>
        <span class="up-body">
          <span class="up-title">添加报告文件</span>
          <span class="up-sub">支持 .xlsx / .pdf / .docx / .txt</span>
        </span>
        <span class="up-btn">选择文件</span>
        <input type="file" id="multifile" accept=".xlsx,.xls,.pdf,.docx,.md,.markdown,.txt,.csv" onchange="onMultiFile(event)">
      </label>
    </div>
    <div class="fields" id="sample-list" style="margin-top:14px"></div>
    <div class="fields" style="margin-top:14px;padding-top:16px;border-top:1px solid var(--line)">
      <div class="field"><span>内标物质</span><input type="text" id="m-ist" value="2-辛醇" oninput="document.getElementById('m-ist').dataset.touched='1';"></div>
      <div class="field"><span>内标浓度</span><input type="text" id="m-cis" value="10"></div>
      <div class="field"><span>浓度单位</span><input type="text" id="m-cisu" value="mg/L"></div>
      <div class="field"><span>内标加量</span><input type="text" id="m-vis" value="50"></div>
      <div class="field"><span>体积单位</span><input type="text" id="m-visu" value="μL"></div>
      <div class="field"><span>样品量</span><input type="text" id="m-ms" value="2"></div>
      <div class="field"><span>取样量单位</span><input type="text" id="m-msu" value="g"></div>
    </div>
    <div class="conc-note">统一内标参数（各样品通用）：C = (内标浓度 × 内标加量 × 响应比 rr) ÷ 样品量 m，结果换算到 μg/L。rr 已含内标峰面积，不再除以内标峰面积。若报告自带浓度列则优先用报告值。</div>
    <div class="row" style="margin-top:14px">
      <button onclick="doMulti()">开始对比分析</button>
      <button class="ghost" onclick="clearMulti()">清空</button>
    </div>
    <div id="multi-out"></div>
  </div>
</div>

<!-- 浏览数据库 -->
<div class="panel" id="p-browse">
  <div class="card">
    <p class="hint">内置风味化合物数据库全量浏览：可按类别筛选、关键词搜索（英文名/中文名），展示中文名、CAS、气味描述与来源。</p>
    <div class="row">
      <select id="bcat" onchange="browseDB()">
        <option value="">全部类别</option>
        <option value="醛类">醛类</option>
        <option value="酮类">酮类</option>
        <option value="酯类">酯类</option>
        <option value="酸类">酸类</option>
        <option value="醇类">醇类</option>
        <option value="萜烯类">萜烯类</option>
        <option value="含硫化合物">含硫化合物</option>
        <option value="吡嗪类">吡嗪类</option>
        <option value="内酯类">内酯类</option>
        <option value="其他">其他</option>
      </select>
      <input type="text" id="bkey" placeholder="搜索：如 丁酸 / Limonene / 奶酪" oninput="browseDB()">
      <span class="hint" id="bcount"></span>
    </div>
    <div id="browse-out"></div>
  </div>
</div>

<script translate="no">
// 防注入/翻译扩展损坏事件的兜底：双绑 + 即使被覆盖也能修复点击
(function(){
  try{
    window.addEventListener('error', function(e){ console.warn('[flavor-tool] captured error', e.message); });
    document.addEventListener('DOMContentLoaded', function(){
      function bindTabs(){
        var tabs=document.querySelectorAll('.tab');
        if(!tabs.length){ setTimeout(bindTabs, 80); return; }
        tabs.forEach(function(t){
          t.onclick=function(){
            document.querySelectorAll('.tab').forEach(function(x){x.classList.remove('active');});
            document.querySelectorAll('.panel').forEach(function(x){x.classList.remove('active');});
            t.classList.add('active');
            var p=document.getElementById('p-'+t.dataset.t); if(p) p.classList.add('active');
          };
        });
      }
      bindTabs();
      var q=document.getElementById('q'); if(q) q.addEventListener('keydown', function(e){ if(e.key==='Enter') doSingle(); });
    });
  }catch(_){ }
})();
const CAT={{ cat_json | safe }};
let lastBatch=[], lastReport=[], lastConc=[];

// 信息完整度：优先展示「有风味描述 + 有阈值 + 有来源」的物质
// 阈值区分度最大（库内约 62% 有值），权重 2；风味描述、来源各 1；满分 4
function hasField(c,k){ const v=c&&c[k]; if(v==null) return false; const s=String(v).trim(); return s!==''&&!['—','-','–','暂无','无','nan','None','null'].includes(s); }
function completeness(c){ if(!c) return 0; return (hasField(c,'odor')?1:0)+(hasField(c,'thr')?2:0)+(hasField(c,'source')?1:0); }
function compLabel(c){ const s=completeness(c); return s>=4?'信息完整':(s>=2?'部分缺失':'信息待补'); }
function sortByCompleteness(rows){ return (rows||[]).slice().sort((a,b)=>completeness(b)-completeness(a)); }

// 内标法浓度: C = (内标浓度 × 内标体积 × 响应比) / 样品取样量
// 响应比 rr = 化合物峰面积/内标峰面积(无量纲)，已含内标峰面积，故不再除以内标峰面积
// 单位换算: 把 (内标浓度单位 × 内标体积单位)/样品量单位 折算到 μg/L（水相近似）
const _MASS={'pg':1e-12,'ng':1e-9,'µg':1e-6,'μg':1e-6,'ug':1e-6,'mg':1e-3,'g':1.0,'kg':1e3};
const _VOL={'pl':1e-12,'nl':1e-9,'µl':1e-3,'µL':1e-3,'μl':1e-3,'μL':1e-3,'ul':1e-3,'UL':1e-3,'mL':1.0,'ml':1.0,'L':1e3,'l':1e3};
function jsParseConcUnit(expr){ if(!expr) return null;
  let s=expr.replace(/[ ×xX()（）]/g,''); if(!s) return null;
  let conv=1;
  for(let tok of s.split('*')){
    let sign=1;
    if(tok.startsWith('/')){ sign=-1; tok=tok.slice(1); }
    if(!tok) continue;
    if(tok.indexOf('/')>=0){ // 复合单位 mass/vol、mass/mass、vol/mass
      const parts=tok.split('/'), num=parts[0], den=parts[1];
      if(_MASS[num]!=null && _VOL[den]!=null) conv*=Math.pow(_MASS[num],sign)*Math.pow(_VOL[den],-sign);
      else if(_MASS[num]!=null && _MASS[den]!=null) conv*=Math.pow(_MASS[num],sign)*Math.pow(_MASS[den],-sign);
      else if(_VOL[num]!=null && _MASS[den]!=null) conv*=Math.pow(_VOL[num],sign)*Math.pow(_MASS[den],-sign);
      else return null;
    } else if(_MASS[tok]!=null) conv*=Math.pow(_MASS[tok],sign);
    else if(_VOL[tok]!=null) conv*=Math.pow(_VOL[tok],sign);
    else return null;
  }
  return conv*1e9; // 折算到 μg/L（水相近似，μg/kg≈μg/L）
}
// conc(): 计算单物质浓度，返回 μg/L(或 null)。直接使用响应比 rr 计算，不再除以内标峰面积
function conc(cis,vis,resp,ms,cisu,visu,msu){
  const a=parseFloat(cis),b=parseFloat(vis),c=parseFloat(resp),d=parseFloat(ms);
  if([a,b,c,d].some(v=>isNaN(v))||d===0) return null;
  const raw=a*b*c/d;
  const unitExpr='('+(cisu||'mg/L')+')*('+(visu||'μL')+')/('+(msu||'g')+')';
  const k=jsParseConcUnit(unitExpr);
  return (k==null)? null : raw*k;
}
// 浓度单位换算系数: 原单位 -> 目标单位（基于 μg/kg = 1 基准）
const _CU={ 'μg/kg':1, 'mg/kg':1e3, 'ng/g':1, 'ng/kg':1e-3, 'g/kg':1e6 };
function concUnitFactor(src,dst){ const a=_CU[src], b=_CU[dst]; if(a==null||b==null) return null; return b/a; }
function concUnitFactor(src,dst){ const a=_CU[src], b=_CU[dst]; if(a==null||b==null) return null; return b/a; }

document.querySelectorAll('.tab').forEach(t=>t.onclick=()=>{
  document.querySelectorAll('.tab').forEach(x=>x.classList.remove('active'));
  document.querySelectorAll('.panel').forEach(x=>x.classList.remove('active'));
  t.classList.add('active');
  document.getElementById('p-'+t.dataset.t).classList.add('active');
});
document.getElementById('q').addEventListener('keydown',e=>{if(e.key==='Enter')doSingle();});

function catCls(cat){return 'b-'+(cat||'其他').replace(/[\\/\\s]/g,'');}
function badge(cat){return `<span class="badge ${catCls(cat)}">${cat||'其他'}</span>`;}
function mcls(m){if(m==='精确')return 'ok';if(m.includes('自动'))return 'bad';return 'warn';}
function srcFmt(s){
  // 来源按行：以「类别：」开头的行，将类别词加粗
  return String(s??'').split('\n').map(l=>{
    const m=l.match(/^([^：:]+)[：:]/);
    return m? `<b>${m[1]}</b>${l.slice(m[1].length)}` : l;
  }).join('<br>');
}
// 仪表盘摘要：保留 响应比/浓度 KPI（报告解析处可关闭 noKpi），并合并「类别覆盖 + 类别分布」为单一可点击模块
function renderSummary(rows, opts){
  opts=opts||{};
  if(!rows||!rows.length) return '';
  const total=rows.length;
  const cats={};rows.forEach(r=>{ if(r.cat) cats[r.cat]=(cats[r.cat]||0)+1; });
  const catCount=Object.keys(cats).length;
  const rrs=rows.filter(r=>typeof r.rr==='number').map(r=>r.rr);
  const sumRR = rrs.reduce((a,b)=>a+b,0);
  const maxRR = rrs.length? Math.max.apply(null,rrs) : null;
  const hasRR = rrs.length>0;
  const concs=rows.filter(r=>typeof r.conc==='number').map(r=>r.conc);
  const hasConc=concs.length>0;
  const maxConc=hasConc? Math.max.apply(null,concs) : null;
  const top = rows.filter(r=>typeof r.rr==='number').slice().sort((a,b)=>b.rr-a.rr).slice(0,3);
  const catArr=Object.entries(cats).sort((a,b)=>b[1]-a[1]);
  const maxCat = catArr.length? catArr[0][1] : 1;
  // 1) KPI 卡片（报告解析处可关闭 noKpi；已删除「命中数」，保留 响应比总和/最高响应比/最高浓度）
  const tag=opts.scope||'本次';
  const kpiHtml = (opts.noKpi ? [] : [
    // 响应比总和：粉
    hasRR?`<div class="kpi pink">
       <div class="kpi-head">
         <span class="kpi-ic">↗</span>
         <div class="kpi-meta"><div class="kpi-tag">相对量纲</div><div class="kpi-trend">∑</div></div>
       </div>
       <div class="kpi-title">响应比总和</div>
       <div class="kpi-value">${sumRR.toFixed(2)}</div>
       <div class="kpi-foot">最高 <b>${maxRR.toFixed(3)}</b> · 来自 <b>${top[0]?top[0].en:'—'}</b></div>
     </div>`:'',
    // 最高响应比：紫
    hasRR?`<div class="kpi lav">
       <div class="kpi-head">
         <span class="kpi-ic">★</span>
         <div class="kpi-meta"><div class="kpi-tag">峰值</div><div class="kpi-trend">#1</div></div>
       </div>
       <div class="kpi-title">最高响应比</div>
       <div class="kpi-value">${maxRR.toFixed(3)}</div>
       <div class="kpi-foot">${top[0]?('<b>'+top[0].en+'</b> '+(top[0].cn||'')):'—'}</div>
     </div>`:'',
    // 浓度峰值：天蓝
    hasConc?`<div class="kpi sky">
       <div class="kpi-head">
         <span class="kpi-ic">μg</span>
         <div class="kpi-meta"><div class="kpi-tag">内标法</div><div class="kpi-trend">max</div></div>
       </div>
       <div class="kpi-title">最高浓度</div>
       <div class="kpi-value">${maxConc.toFixed(3)}<span class="unit">μg/L</span></div>
       <div class="kpi-foot">共 <b>${concs.length}</b> 项参与计算</div>
     </div>`:''
  ]).filter(Boolean).join('');
  // 2) 合并模块：类别覆盖（每类一张可点击图标卡，按数量排序；点击 → 下方列出该类别物质）
  const catDot=cat=>`<span class="cc-dot ${catCls(cat)}" aria-hidden="true"></span>`;
  let catHtml='<div class="cat-bars"><h3><span class="ic">▤</span>类别覆盖<span class="sub">（共 '+catCount+' 类 · 点击类别展开物质）</span></h3><div class="cat-grid">';
  catArr.forEach(([k,v])=>{
    catHtml += `<div class="cat-chip cat-row" data-cat="${k}" onclick="filterSummaryCat(this)" title="点击查看「${k}」全部物质">
        <span class="cc-head">${catDot(k)}<span class="cc-name">${k}</span></span>
        <span class="cc-num">${v}<small>种</small></span></div>`;
  });
  catHtml += '</div></div>';
  // 3) Top 3
  let topHtml='';
  if(top.length){
    topHtml = '<div class="top-list"><h3><span class="ic">★</span>响应比 Top 3</h3>';
    top.forEach((r,i)=>{
      const v = typeof r.rr==='number'? r.rr.toFixed(3) : '—';
      topHtml += `<div class="top-item"><div class="top-rank">${i+1}</div><div class="top-name"><div class="en">${r.en}</div><div class="cn">${badge(r.cat||'其他')} ${r.cn||''}</div></div><div class="top-val">${v}<small>rr</small></div></div>`;
    });
    topHtml += '</div>';
  }
  const kpiBlock = kpiHtml? `<div class="kpis">${kpiHtml}</div>` : '';
  const catConc = renderCatConc(rows, opts);
  return `${kpiBlock}<div style="display:grid;grid-template-columns:1.2fr 1fr;gap:14px">${catHtml}${topHtml}</div>`+
    catConc +
    `<div class="cat-detail" id="cat-detail" style="margin-top:14px"></div>`;
}

// 分类浓度总量：把每个物质的内标法浓度(μg/L)按类别汇总，展示占比条形 + 合计
function renderCatConc(rows, opts){
  opts = opts || {};
  if(!rows || !rows.length) return '';
  const tot = {}; let grand = 0, nConc = 0;
  rows.forEach(r=>{
    const c = r.cat || '其他';
    const v = (typeof r.conc === 'number' && !isNaN(r.conc)) ? r.conc : 0;
    tot[c] = (tot[c] || 0) + v;
    if(v > 0) nConc++;
    grand += v;
  });
  const arr = Object.entries(tot).filter(([k,v])=>v > 0).sort((a,b)=>b[1]-a[1]);
  if(!arr.length) return '';
  const max = arr[0][1] || 1;
  const unit = opts.unit || 'μg/L';
  let h = '<div class="card" style="margin-top:16px"><h3><span class="ic">⚖</span>分类浓度总量'+
          '<span class="sub">（内标法浓度按类别汇总 · '+unit+'）</span></h3><div class="conc-rows">';
  arr.forEach(([k,v])=>{
    const pct = grand ? (v/grand*100) : 0;
    const w = max ? (v/max*100) : 0;
    h += `<div class="conc-row"><span class="cb">${badge(k)}</span>`+
         `<span class="ct"><span class="cf" style="width:${w.toFixed(1)}%"></span></span>`+
         `<span class="cv">${v.toFixed(2)}<small>${unit}</small></span>`+
         `<span class="cp">${pct.toFixed(1)}%</span></div>`;
  });
  h += `</div><div class="conc-foot"><span>合计浓度 <b>${grand.toFixed(2)} ${unit}</b></span>`+
       `<span>参与计算 <b>${nConc}</b> 项</span><span>覆盖 <b>${arr.length}</b> 类</span></div></div>`;
  return h;
}
// 点击类别 → 在下方列出该类别的物质（作用于最近一次 renderSummary 的 rows）
let _lastSummaryRows=null;
function _setSummaryRows(rows){ _lastSummaryRows=rows; }
function filterSummaryCat(el){
  const cat=el.getAttribute('data-cat'); if(!cat) return;
  const box=document.getElementById('cat-detail');
  if(!box) return;
  // 高亮当前卡片
  document.querySelectorAll('.cat-row').forEach(c=>c.classList.remove('is-active'));
  if(box._cat===cat){ box.innerHTML=''; box._cat=null; return; }  // 再次点击收起
  el.classList.add('is-active');
  box._cat=cat;
  // 排序：先按信息完整度（有风味描述/阈值/来源优先），再按 OAV 由大到小
  const rows=(_lastSummaryRows||[]).filter(r=>r.cat===cat).slice().sort((a,b)=>{
    const dc=completeness(b)-completeness(a);
    if(dc!==0) return dc;
    const av=(r)=> (typeof r.oav==='number'&&!isNaN(r.oav))?r.oav:0;
    return av(b)-av(a);
  });
  const flsOf=r=>{ const f=r.oav_flag||'—'; return f==='关键致香'?'f-key':(f==='潜在贡献'?'f-pot':'f-na'); };
  let h='<div class="card" style="margin:0"><h3><span class="ic">▤</span>'+cat+' · 共 '+rows.length+' 种物质（按 OAV 由大到小）</h3><div class="wrap"><table><thead><tr><th>#</th><th>英文名</th><th>中文名</th><th>类别</th><th>阈值(μg/L)</th><th>浓度(μg/L)</th><th>OAV</th><th>ROAV</th><th>气味活性</th><th>气味描述</th></tr></thead><tbody>';
  h+=rows.map((r,i)=>{
    const cu=(r.conc==null||isNaN(r.conc))?'—':Number(r.conc).toPrecision(4).replace(/\.?0+$/,'');
    const oav=(r.oav==null)?'—':r.oav.toPrecision(4).replace(/\.?0+$/,'');
    const roav=(r.roav==null)?'—':r.roav.toFixed(2);
    const flag=r.oav_flag||'—'; const fls=flsOf(r);
    return `<tr><td>${i+1}</td><td><b>${r.en}</b></td><td class="cn">${r.cn??''}</td><td>${badge(r.cat)}</td><td>${r.thr}</td><td class="num-dark">${cu}</td><td class="num-dark">${oav}</td><td class="num-dark">${roav}</td><td class="oavflag ${fls}">${flag}</td><td>${r.odor??''}</td></tr>`;
  }).join('');
  h+='</tbody></table></div></div>';
  box.innerHTML=h;
}
function tableHTML(rows){
  if(!rows.length)return '<p class="hint">无结果</p>';
  rows=sortByCompleteness(rows);  // 信息完整（有气味描述/阈值/来源）的优先展示
  let h='<div class="wrap"><table><thead><tr><th>#</th><th>英文名</th><th>中文名</th><th>CAS</th><th>类别</th><th>阈值(μg/L)</th><th>介质</th><th>匹配</th><th>气味描述</th><th>来源</th></tr></thead><tbody>';
  rows.forEach((r,i)=>{
    h+=`<tr><td>${i+1}</td><td><b>${r.en}</b></td><td class="cn">${r.cn??''}</td><td>${r.cas??''}</td><td>${badge(r.cat)}</td><td>${r.thr}</td><td>${r.med}</td><td class="m ${mcls(r.match)}">${r.match}</td><td class="desc">${r.odor??''}</td><td class="src">${srcFmt(r.source)}</td></tr>`;
  });
  h+='</tbody></table></div>';
  return h;
}
function statHTML(rows){
  const c={};rows.forEach(r=>c[r.cat]=(c[r.cat]||0)+1);
  return '<div class="stat">'+Object.entries(c).map(([k,v])=>`<span>${badge(k)} <b>${v}</b></span>`).join('')+'</div>';
}

async function doSingle(){
  const q=document.getElementById('q').value.trim();if(!q)return;
  const r=await(await fetch('/api/search',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({names:[q]})})).json();
  _setSummaryRows(r);
  document.getElementById('single-out').innerHTML=renderSummary(r,{scope:'单物质查询'})+tableHTML(r);
}
async function doBatch(){
  const txt=document.getElementById('batch').value;
  const names=txt.split(/\n|\r/).map(s=>s.trim()).filter(Boolean);
  if(!names.length)return;
  lastBatch=await(await fetch('/api/search',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({names})})).json();
  _setSummaryRows(lastBatch);
  document.getElementById('batch-out').innerHTML=renderSummary(lastBatch,{scope:'批量检索'})+tableHTML(lastBatch);
}
async function loadBatchFile(){
  const el=document.getElementById('batchfile');
  const f=el.files[0];if(!f)return;
  const fd=new FormData();fd.append('file',f);
  document.getElementById('batch').value='解析中…';
  try{
    const r=await(await fetch('/api/batch_upload',{method:'POST',body:fd})).json();
    document.getElementById('batch').value=(r.names||[]).join('\n');
    if((r.names||[]).length===0){
      document.getElementById('batch').value='';
      alert('未在文件中识别到已知风味物质，请确认文件内容或手动粘贴物质名。');
    }
  }catch(e){
    document.getElementById('batch').value='';
    alert('文件解析失败：'+e.message);
  }
}
async function doReport(){
  const f=document.getElementById('reportfile').files[0];if(!f)return;
  const out=document.getElementById('report-out');
  out.innerHTML='<p class="hint" style="color:var(--teal)">⏳ 解析中…正在读取并识别报告（化合物 / 内标 / 响应比）…</p>';
  const fd=new FormData();fd.append('file',f);fd.append('ist_name',document.getElementById('r-ist').value.trim());
  let resp;
  try{
    resp=await(await fetch('/api/upload',{method:'POST',body:fd})).json();
  }catch(e){
    // 网络/解析异常（如后端返回非 JSON 错误页）不再静默失败
    out.innerHTML='<p class="hint" style="color:#e05a4a">报告上传或解析失败：'+escapeHtml(e.message)+'</p>';
    return;
  }
  // 兼容后端错误响应 {error, rows}
  if(resp && resp.error){
    out.innerHTML='<p class="hint" style="color:#e05a4a">'+escapeHtml(resp.error)+'</p>';
    return;
  }
  // 多处理组报告：自动识别处理组并横向对比
  if(resp && resp.multi){
    lastMulti=resp; lastReport=[];
    const ie=document.getElementById('r-ist'); if(ie) ie.value=resp.ist_name||ie.value;
    recalcGroupCompare();
    return;
  }
  const rows=(resp && resp.rows) ? resp.rows : (Array.isArray(resp)?resp:[]);
  if(!rows.length){
    out.innerHTML='<p class="hint">未从报告中识别到已知风味物质。请确认文件为 GC-MS 定量报告（含「化合物/英文名 + 响应比/峰面积」等表头），或文献/报告文档。</p>';
    lastReport=[]; return;
  }
  // 用后端自动识别的内标物质名预填（用户可手动改）
  if(resp.ist_name){ const ie=document.getElementById('r-ist'); if(ie && !ie.dataset.touched) ie.value=resp.ist_name; }
  lastReport=rows;
  // 保留报告自带「最终浓度」列，供未填内标参数时回退显示
  lastReport.forEach(r=>{
    r._conc_raw = (typeof r.conc==='number') ? r.conc : (typeof r.conc==='string' ? parseFloat(r.conc) : null);
    r._conc_unit_raw = (typeof r.conc_unit==='string') ? r.conc_unit : '';
  });
  await applyReportConc();
  renderReport();
}
async function applyReportConc(){
  const cis=document.getElementById('r-cis').value.trim();
  const vis=document.getElementById('r-vis').value.trim();
  const ms=document.getElementById('r-ms').value.trim();
  const cisu=document.getElementById('r-cisu').value.trim()||'mg/L';
  const visu=document.getElementById('r-visu').value.trim()||'μL';
  const msu=document.getElementById('r-msu').value.trim()||'g';
  const ok = cis!==''&&vis!==''&&ms!==''&&!isNaN(parseFloat(cis))&&!isNaN(parseFloat(vis))&&!isNaN(parseFloat(ms))&&parseFloat(ms)!==0;
  lastReport.forEach(r=>{
    // 优先用报告自带浓度列(若含有效浓度列且未给内标参数则直接展示)；给了内标参数则按公式算
    if(ok){
      const resp = (typeof r.rr==='number' && r.rr!=null) ? r.rr : null;
      const c = resp!=null ? conc(cis,vis,resp,ms,cisu,visu,msu) : null;
      r.conc = c; r.conc_unit='μg/L'; r.conc_ugkg=c;
    } else if(typeof r.conc0==='number' && r.conc0!=null){
      r.conc = r.conc0; r.conc_unit=(r.conc_unit||'μg/L'); r.conc_ugkg=r.conc0;
    } else { r.conc=null; r.conc_unit=null; r.conc_ugkg=null; }
  });
  lastReport._conc_ok = ok;
  lastReport._conc_unit = ok ? 'μg/L' : null;
  computeOAV(lastReport);
}
// OAV(气味活性值)=C/阈值; ROAV(相对气味活性值)=OAV_i/max(OAV)×100
// 浓度优先用折算到 μg/kg 的 conc_ugkg，回退用响应比 rr(相对量纲)
function parseThr(s){
  if(!s || typeof s!=='string') return null;
  let t=s.replace(/[（(].*$/,'') .replace(/[≈~]/g,'').replace(/[ \u00d7]/g,'').replace(/[—\-—]/g,'');
  // 科学计数 1.0×10⁵
  let m=t.match(/([0-9.]+)[x×]10\^?([0-9]+)/);
  if(m) return parseFloat(m[1])*Math.pow(10,parseInt(m[2]));
  if(t.indexOf('–')>=0||t.indexOf('-')>=0){
    let ps=t.split(/[–-]/).map(parseFloat).filter(x=>!isNaN(x));
    if(ps.length) return ps.reduce((a,b)=>a+b,0)/ps.length;
  }
  let v=parseFloat(t);
  return isNaN(v)?null:v;
}
function computeOAV(rows){
  const pairs=[];
  rows.forEach(r=>{
    const T=parseThr(r.thr);
    if(T==null||!(T>0)) { r.oav=null; r.roav=null; r.oav_flag='—'; return; }
    let val=null;
    if(typeof r.conc_ugkg==='number' && r.conc_ugkg!=null) val=r.conc_ugkg;
    else if(typeof r.conc==='number' && r.conc!=null) val=r.conc;
    else if(typeof r.rr==='number' && r.rr!=null) val=r.rr;
    if(val==null){ r.oav=null; r.roav=null; r.oav_flag='—'; return; }
    r.oav = val/T;
    pairs.push(r);
  });
  const mx = pairs.length? Math.max.apply(null, pairs.map(r=>r.oav)) : 0;
  pairs.forEach(r=>{
    r.roav = mx? r.oav/mx*100 : 0;
    if(r.roav>=10) r.oav_flag='关键致香';
    else if(r.roav>=1) r.oav_flag='潜在贡献';
    else r.oav_flag='—';
  });
}
function renderReport(){
  if(!lastReport.length){document.getElementById('report-out').innerHTML='<p class="hint">无结果</p>';return;}
  const dst=document.getElementById('r-cu-dst') ? document.getElementById('r-cu-dst').value : 'μg/L';
  let h='';
  const istA=(lastReport&&lastReport.istd_area!=null)?lastReport.istd_area:(lastReport&&lastReport[0]&&lastReport[0].istd_area!=null?lastReport[0].istd_area:null);
  if(istA!=null){ const istName=document.getElementById('r-ist').value.trim()||'内标'; h+='<p class="hint" style="color:var(--teal)">已识别内标物质「'+istName+'」的峰面积 A₁ = '+istA+'（已包含在响应比 rr = A÷A₁ 中，公式不再单独除以 A₁）。</p>'; }
  h+=renderSummary(lastReport,{scope:'GC-MS 报告',noKpi:true});
  _setSummaryRows(lastReport);
  // 类别筛选（下拉单选，与多报告模块 m-cat-filter 风格一致）
  const cats=[...new Set(lastReport.map(r=>r.cat||'其他'))];
  const catOrder=(typeof CAT!=='undefined'&&CAT)?(Array.isArray(CAT)?CAT.map(c=>c.name||c):Object.keys(CAT)):[];
  cats.sort((a,b)=>{var ia=catOrder.indexOf(a),ib=catOrder.indexOf(b);ia=ia<0?999:ia;ib=ib<0?999:ib;return ia-ib;});
  if(reportCatFilter && cats.indexOf(reportCatFilter)<0) reportCatFilter='';  // 数据变更后失效的筛选自动清空
  const fsel=reportCatFilter||'';
  let fbar='<div class="gc-catbar" style="margin:8px 0 4px"><label>类别筛选：<select id="r-cat-filter" onchange="reportCatFilter=this.value;renderReport();"><option value="">全部类别</option>';
  cats.forEach(c=>{ fbar+=`<option value="${escapeHtml(c)}"${c===fsel?' selected':''}>${escapeHtml(c)}</option>`; });
  fbar+='</select></label><span class="hint">按类别筛选主数据表物质</span></div>';
  h+=fbar;
  const rows=lastReport.filter(r=> !reportCatFilter || (r.cat||'其他')===reportCatFilter);
  h+='<div class="wrap"><table id="report-table"><thead><tr><th>#</th><th>英文名</th><th>中文名</th><th>CAS</th><th>类别</th><th>RT(min)</th><th>响应比 rr</th><th>阈值(μg/L)</th><th>浓度('+dst+')</th><th>OAV</th><th>ROAV</th><th>气味活性</th><th>气味描述</th><th>来源</th></tr></thead><tbody>';
  rows.forEach((r,i)=>{
    const cu = (r.conc==null||isNaN(r.conc)) ? '—' : Number(r.conc).toPrecision(4).replace(/\.?0+$/,'');
    const oav = (r.oav==null) ? '—' : r.oav.toPrecision(4).replace(/\.?0+$/,'');
    const roav = (r.roav==null) ? '—' : r.roav.toFixed(2);
    const flag = r.oav_flag||'—';
    const fls = flag==='关键致香'?'f-key':(flag==='潜在贡献'?'f-pot':'f-na');
    h+=`<tr><td>${i+1}</td><td><b>${r.en}</b></td><td class="cn">${r.cn??''}</td><td>${r.cas??''}</td><td>${badge(r.cat)}</td><td>${r.rt??''}</td><td>${r.rr??''}</td><td>${r.thr}</td><td>${cu}</td><td>${oav}</td><td>${roav}</td><td class="oavflag ${fls}">${flag}</td><td class="desc">${r.odor??''}</td><td class="src">${srcFmt(r.source)}</td></tr>`;
  });
  h+='</tbody></table></div>';
  document.getElementById('report-out').innerHTML=h;
}
function recalcReportConc(){
  if(!lastReport.length){alert('请先解析报告');return;}
  applyReportConc().then(renderReport);
}
function liveRecalcReport(){
  if(lastReport && lastReport.length){ applyReportConc().then(renderReport); }
  if(lastMulti && lastMulti.compounds){ recalcGroupCompare(); }
}

/* ===================== 处理组对比（单 PDF 多处理组自动识别） ===================== */
let lastMulti=null;   // {multi,groups,samples,ist_name,ist_area_by_group,compounds,rows}
const GC_COLORS=['#24908c','#e0844a','#4a7de0','#b0556f','#7a9b3b'];
function recalcGroupCompare(){
  if(!lastMulti || !lastMulti.compounds) return;
  const cis=document.getElementById('r-cis').value.trim();
  const vis=document.getElementById('r-vis').value.trim();
  const ms=document.getElementById('r-ms').value.trim();
  const cisu=document.getElementById('r-cisu').value.trim()||'mg/L';
  const visu=document.getElementById('r-visu').value.trim()||'μL';
  const msu=document.getElementById('r-msu').value.trim()||'g';
  const ok = cis!==''&&vis!==''&&ms!==''&&!isNaN(parseFloat(cis))&&!isNaN(parseFloat(vis))&&!isNaN(parseFloat(ms))&&parseFloat(ms)!==0;
  const groups=lastMulti.groups;
  lastMulti.compounds.forEach(c=>{
    const T=parseThr(c.thr);
    c._T=T;
    groups.forEach(g=>{
      const cell=c.by_group[g];
      if(!cell){ return; }
      let C=null,oav=null;
      if(ok && typeof cell.rr==='number' && cell.rr!=null) C=conc(cis,vis,cell.rr,ms,cisu,visu,msu);
      if(C!=null && T && T>0) oav=C/T;
      cell._conc=C; cell._oav=oav;
    });
  });
  lastMulti._conc_ok=ok;
  renderGroupCompare();
}
function renderGroupCompare(){
  const box=document.getElementById('report-out');
  if(!lastMulti || !lastMulti.compounds){ box.innerHTML='<p class="hint">无处理组数据</p>'; gcCatFilter=null; return; }
  // 重新渲染时保留当前类别筛选（单选下拉），未选则显示全部
  const gcSelPrev=(document.getElementById('gc-cat-filter-sel')||{}).value||'';
  gcCatFilter = gcSelPrev ? new Set([gcSelPrev]) : null;
  const groups=lastMulti.groups, samples=lastMulti.samples;
  const ctrl = groups.find(g=>/ck|对照|control|空白/i.test(g)) || groups[groups.length-1];
  const istArea = lastMulti.ist_area_by_group||{};
  let chips = groups.map((g,i)=>`<span class="gc-chip"><span class="dot" style="background:${GC_COLORS[i%GC_COLORS.length]}"></span>${escapeHtml(g)} <span style="color:var(--mut-2);font-weight:600">(${escapeHtml(samples[i]||'')})</span></span>`).join('');
  let h=`<div class="gc-banner"><div>🧪 已自动识别 <b>${groups.length}</b> 个处理组，并提取各组物质定量数据（内标法：浓度 = 内标浓度 × 内标加量 × 响应比 rr ÷ 样品量）。</div>`;
  h+=`<div class="gc-chips">${chips}</div>`;
  h+=`<div class="gc-chips" style="margin-top:6px"><span class="gc-chip" style="background:var(--card-soft)">内标：<b>${escapeHtml(lastMulti.ist_name||'—')}</b></span>`;
  groups.forEach(g=>{ if(istArea[g]!=null) h+=`<span class="gc-chip" style="background:var(--card-soft)">A₁(${escapeHtml(g)})=${fmtNum(istArea[g])}</span>`; });
  h+=`</div></div>`;
  // 摘要统计
  const presentCount={}; groups.forEach(g=>presentCount[g]=0);
  const shared=[]; const uniqueByG={}; groups.forEach(g=>uniqueByG[g]=[]);
  const upByG={}; const downByG={}; groups.forEach(g=>{upByG[g]=[];downByG[g]=[];});
  lastMulti.compounds.forEach(c=>{
    const pres=groups.filter(g=>c.by_group[g]);
    pres.forEach(g=>presentCount[g]++);
    if(pres.length===groups.length) shared.push(c);
    else if(pres.length===1) pres.forEach(g=>uniqueByG[g].push(c));
    if(c.by_group[ctrl] && typeof c.by_group[ctrl].rr==='number' && c.by_group[ctrl].rr>0){
      groups.forEach(g=>{ if(g===ctrl) return; const cell=c.by_group[g];
        if(cell && typeof cell.rr==='number' && cell.rr>0){ const fc=cell.rr/c.by_group[ctrl].rr;
          if(fc>=2) upByG[g].push({c,fc}); else if(fc<=0.5) downByG[g].push({c,fc}); } });
    }
  });
  h+=`<div class="gc-summary">`;
  h+=statCard('共有物质', shared.length, '在全部 '+groups.length+' 个处理组中均检出');
  groups.forEach(g=> h+=statCard('特有物质 · '+g, uniqueByG[g].length, uniqueByG[g].map(c=>dispName(c)).slice(0,4).join('、')||'—'));
  groups.forEach(g=>{ if(g!==ctrl) h+=statCard('上调(≥2×) · '+g, upByG[g].length, '相对对照 '+escapeHtml(ctrl)); });
  groups.forEach(g=>{ if(g!==ctrl) h+=statCard('下调(≤0.5×) · '+g, downByG[g].length, '相对对照 '+escapeHtml(ctrl)); });
  h+=`</div>`;
  h+=`<div class="gc-tools"><button class="ghost" onclick="exportGroupCSV()">⬇ 导出对比矩阵(CSV)</button><button class="ghost" onclick="exportGroupCatCSV()">⬇ 导出类别浓度总量(CSV)</button><span class="hint" style="align-self:center">内标参数变动将实时重算浓度 / OAV（见上方输入框）。</span></div>`;
  // 管理工具条（删除 / 保存 / 撤回）
  h+=`<div class="gc-mgrbar"><b>主表物质管理</b>`
    +`<button class="ghost" id="gc-undo-single">撤回删除</button>`
    +`<button class="ghost" id="gc-save-single">保存调整</button>`
    +`<span class="hint" id="gc-stat-single"></span>`
    +`<span class="hint">点「✕」即删除该物质（与主可视化面板共享删除状态，可「撤回删除」并「保存调整」到浏览器本地）。导出将剔除已删除物质。</span></div>`;
  // 类别筛选（下拉单选，与 report-table 的 r-cat-filter / 多报告的 m-cat-filter 风格一致）
  const gcCats=[...new Set(lastMulti.compounds.map(function(c){return c.cat||'其他';}))];
  const gcOrder=(typeof CAT!=='undefined'&&CAT)?(Array.isArray(CAT)?CAT.map(function(c){return c.name||c;}):Object.keys(CAT)):[];
  gcCats.sort(function(a,b){var ia=gcOrder.indexOf(a),ib=gcOrder.indexOf(b);ia=ia<0?999:ia;ib=ib<0?999:ib;return ia-ib;});
  let gopt='<option value="">全部类别</option>';
  gcCats.forEach(function(c){ gopt+='<option value="'+escapeHtml(c)+'"'+(c===gcSelPrev?' selected':'')+'>'+escapeHtml(c)+'</option>'; });
  h+=`<div class="gc-catbar" style="margin:8px 0 4px"><label>类别筛选：<select id="gc-cat-filter-sel" onchange="gcSetCatFilter(this.value)">${gopt}</select></label><span class="hint">按类别筛选主对比表物质</span></div>`;
  // 矩阵
  h+=`<div class="wrap"><table id="gc-matrix"><thead><tr><th class="gc-del-col">✕</th><th>物质</th><th>中文名</th><th>类别</th>`;
  groups.forEach(g=> h+=`<th class="grp-col">${escapeHtml(g)}<br><small>响应比 / 浓度 / OAV</small></th>`);
  h+=`<th>相对对照<br><small>${escapeHtml(ctrl)} (rr比)</small></th><th>风味描述</th></tr></thead><tbody id="gc-matrix-body">`;
  h+=gcMatrixRowsHTML();
  h+=`</tbody></table></div>`;
  h+=renderGroupCatConc();
  h+=`<div class="viz-wrap" id="viz-single"></div>`;
  box.innerHTML=h;
  if(window.NatureViz) NatureViz.build('single', lastMulti);
  gcWireMatrix();
}
function maxRR(c){ let m=0; for(const g in c.by_group){ const v=c.by_group[g]&&c.by_group[g].rr; if(typeof v==='number'&&v>m)m=v; } return m; }
function dispName(c){ return (c.cn && c.cn!='(未收录)')? c.cn : c.en; }

// —— 主对比矩阵：直接删除键（隐藏）/ 保存 / 撤回（与可视化面板共享 gcms_nv_single.deleted，导出亦剔除）——
var gcCatFilter=null;  // 主表类别筛选（Set<cat>）；null=未初始化
var reportCatFilter='';  // GC-MS 主数据表（report-table）类别筛选；''=全部
function gcDelSet(kind){
  if(window.NatureViz && NatureViz.state && NatureViz.state[kind]) return NatureViz.state[kind].deleted;
  try{ var raw=localStorage.getItem('gcms_nv_'+kind); var o=raw?JSON.parse(raw):null; return new Set(o&&o.deleted?o.deleted:[]); }catch(e){ return new Set(); }
}
function gcSetCatFilter(v){
  // 主对比表类别筛选（单选下拉）：''=全部类别，否则仅显示该类别
  gcCatFilter = v ? new Set([v]) : null;
  gcRefreshMatrix('single');
}
function gcMatrixRowsHTML(){
  if(!lastMulti) return '';
  const groups=lastMulti.groups, ctrl=groups.find(g=>/ck|对照|control|空白/i.test(g))||groups[groups.length-1];
  const del=gcDelSet('single');
  const comps = (window.NatureViz && NatureViz.sortByCatValueDesc)
    ? NatureViz.sortByCatValueDesc(lastMulti.compounds, c=>maxRR(c))
    : [...lastMulti.compounds].sort((a,b)=>maxRR(b)-maxRR(a));
  return comps.map(function(c){
    const en=c.en;
    if(del.has(en)) return '';                       // 已删除：直接从主表隐藏
    if(gcCatFilter && !gcCatFilter.has(c.cat||'其他')) return '';  // 类别筛选
    const name=dispName(c);
    let row='<tr><td class="gc-del-col"><button class="gc-del-btn" data-en="'+escapeHtml(en)+'" title="删除该物质">✕</button></td>';
    row+='<td><b>'+escapeHtml(en)+'</b></td><td class="cn">'+escapeHtml(name)+'</td><td>'+badge(c.cat)+'</td>';
    let bestG=null,bestRR=-1;
    groups.forEach(g=>{ const cell=c.by_group[g]; if(cell&&typeof cell.rr==='number'&&cell.rr>bestRR){bestRR=cell.rr;bestG=g;} });
    groups.forEach(g=>{
      const cell=c.by_group[g];
      if(!cell){ row+='<td class="gc-na grp-col">—</td>'; return; }
      const rr=cell.rr, C=cell._conc, oav=cell._oav;
      const cls=(g===bestG)?'hl-max':'';
      const rrS=(typeof rr==='number')? rr.toPrecision(3).replace(/\.?0+$/,'') : '—';
      const cS=fmtConc(C), oS=fmtOAV(oav);
      row+='<td class="grp-col '+cls+'"><div class="cell-rr">'+rrS+'</div><div class="cell-conc">'+cS+' μg/L</div><div class="cell-oav '+oavFlagCls(oav)+'">OAV '+oS+'</div></td>';
    });
    const cc=c.by_group[ctrl];
    if(cc && typeof cc.rr==='number' && cc.rr>0){
      let parts=[];
      groups.forEach(g=>{ if(g===ctrl) return; const cell=c.by_group[g];
        if(cell && typeof cell.rr==='number' && cell.rr>0){ const fc=cell.rr/cc.rr;
          const cls=fc>=2?'fc-up':(fc<=0.5?'fc-down':''); const arr=fc>=2?'▲':(fc<=0.5?'▼':'');
          parts.push('<span class="'+cls+'">'+escapeHtml(g)+' '+arr+fc.toFixed(1)+'×</span>'); } });
      row+='<td style="text-align:left">'+(parts.join('<br>')||'—')+'</td>';
    } else row+='<td class="gc-na">—</td>';
    row+='<td class="gc-desc" title="'+escapeHtml(c.odor||'')+'">'+escapeHtml(c.odor||'—')+'</td></tr>';
    return row;
  }).join('');
}
function gcUpdateStat(kind){
  const el=document.getElementById('gc-stat-'+kind); if(!el) return;
  const del=gcDelSet(kind), total=(lastMulti?lastMulti.compounds.length:0);
  el.textContent='已保留 '+(total-del.size)+' / 共 '+total;
}
function gcRefreshMatrix(kind){
  if(kind!=='single') return;
  const tb=document.getElementById('gc-matrix-body'); if(!tb) return;
  tb.innerHTML=gcMatrixRowsHTML(); gcUpdateStat(kind);
}
function gcDeleteOne(kind, en){
  // 直接删除（加入共享删除集并隐藏该行）；撤回删除可恢复
  if(window.NatureViz && NatureViz.state && NatureViz.state[kind]){
    const st=NatureViz.state[kind];
    if(!st.deleted.has(en)){ st.deleted.add(en); st.undoStack.push(en); }
    NatureViz.save(kind); NatureViz.render(kind);
  }
  gcRefreshMatrix(kind);
}
function gcUndo(kind){
  if(window.NatureViz && NatureViz.state && NatureViz.state[kind]){
    const en=NatureViz.state[kind].undoStack.pop();
    if(en){ NatureViz.state[kind].deleted.delete(en); NatureViz.save(kind); NatureViz.render(kind); }
  }
  gcRefreshMatrix(kind);
}
function gcSave(kind){
  if(window.NatureViz && NatureViz.state && NatureViz.state[kind]) NatureViz.save(kind);
  const b=document.getElementById('gc-save-'+kind); if(b){ b.textContent='已保存✓'; setTimeout(function(){ b.textContent='保存调整'; }, 1200); }
}
function gcWireMatrix(){
  const tbl=document.getElementById('gc-matrix'); if(!tbl) return;
  tbl.addEventListener('click', function(e){ const b=e.target.closest('.gc-del-btn'); if(b){ gcDeleteOne('single', b.getAttribute('data-en')); } });
  const undo=document.getElementById('gc-undo-single'); if(undo) undo.addEventListener('click', function(){ gcUndo('single'); });
  const save=document.getElementById('gc-save-single'); if(save) save.addEventListener('click', function(){ gcSave('single'); });
  gcUpdateStat('single');
}
// —— 多报告对比主表：每行删除（小圆点 ✕），与物质管理面板共享删除集（可「撤回删除」/「保存调整」）——
function mcDelSet(){
  if(window.NatureViz && NatureViz.state && NatureViz.state['multi']) return NatureViz.state['multi'].deleted;
  return new Set();
}
function mcRefreshMatrix(){
  const sel=document.getElementById('m-cat-filter'); const f=sel?sel.value:'';
  const box=document.getElementById('multi-matrix-rows'); if(box && window.__buildMultiRows) box.innerHTML=window.__buildMultiRows(f);
}
function mcDeleteMulti(en){
  if(window.NatureViz && NatureViz.state && NatureViz.state['multi']){
    const st=NatureViz.state['multi'];
    if(!st.deleted.has(en)){ st.deleted.add(en); st.undoStack.push(en); }
    NatureViz.save('multi'); NatureViz.render('multi');   // 同步刷新热力图 + 物质管理 + 主表
  } else { mcRefreshMatrix(); }
}
function statCard(title,val,sub){ return `<div class="gc-stat"><h4>${escapeHtml(title)}</h4><div class="v">${val}</div><div class="sub">${escapeHtml(sub||'')}</div></div>`; }
function oavFlagCls(oav){ if(oav==null) return ''; if(oav>=10) return 'f-key'; if(oav>=1) return 'f-pot'; return 'f-na'; }
function fmtNum(x){ if(x==null) return '—'; if(Math.abs(x)>=1e6) return (x/1e6).toFixed(2)+'M'; if(Math.abs(x)>=1e3) return (x/1e3).toFixed(1)+'k'; return ''+Math.round(x); }
function fmtConc(x){ if(x==null||!isFinite(x)) return '—'; if(x>=1e6) return (x/1e6).toFixed(2)+'M'; if(x>=1e4) return (x/1e3).toFixed(1)+'k'; if(x>=100) return x.toFixed(0); return Number(x).toPrecision(3).replace(/\.?0+$/,''); }
function fmtOAV(x){ if(x==null||!isFinite(x)) return '—'; if(x>=1e6) return (x/1e6).toFixed(2)+'M'; if(x>=1e3) return (x/1e3).toFixed(1)+'k'; if(x>=10) return x.toFixed(0); return x.toPrecision(3).replace(/\.?0+$/,''); }
function renderGroupCatConc(){
  const groups=lastMulti.groups;
  const catTotals={}; groups.forEach(g=>catTotals[g]={}); const catSet=[];
  lastMulti.compounds.forEach(c=>{ groups.forEach(g=>{ const cell=c.by_group[g];
    if(cell && cell._conc!=null){ const cat=c.cat||'其他'; if(!catTotals[g][cat])catTotals[g][cat]=0; catTotals[g][cat]+=cell._conc; if(catSet.indexOf(cat)<0)catSet.push(cat); } }); });
  catSet.sort((a,b)=> sumArr(catTotals,a)-sumArr(catTotals,b));
  const grand={}; groups.forEach(g=> grand[g]=catSet.reduce((s,cat)=>s+(catTotals[g][cat]||0),0));
  let h=`<div class="card" style="margin-top:16px"><h3><span class="ic">⚖</span>各处理组类别浓度总量对比 <small style="color:var(--mut)">（内标法浓度 μg/L 按类别汇总）</small></h3>`;
  h+=`<table class="multi-cat"><thead><tr><th>类别</th>`+groups.map(g=>`<th>${escapeHtml(g)}</th>`).join('')+`<th>合计</th></tr></thead><tbody>`;
  catSet.forEach(cat=>{
    h+=`<tr><td>${badge(cat)}</td>`; let rowTot=0;
    groups.forEach(g=>{ const v=catTotals[g][cat]||0; rowTot+=v; const mx=grand[g]||1;
      h+=`<td class="cc-cell"><div class="cc-bar" style="width:${Math.min(100,v/mx*100)}%"></div><span class="cc-v">${v.toFixed(1)}</span></td>`; });
    h+=`<td class="cc-sum">${rowTot.toFixed(1)}</td></tr>`;
  });
  h+=`<tr class="cc-total"><td>合计</td>`+groups.map(g=>`<td>${grand[g].toFixed(1)}</td>`).join('')+`<td>${groups.reduce((s,g)=>s+grand[g],0).toFixed(1)}</td></tr>`;
  h+=`</tbody></table></div>`;
  return h;
}
function sumArr(catTotals,cat){ let s=0; for(const g in catTotals) s+=(catTotals[g][cat]||0); return s; }
function exportGroupCSV(){
  if(!lastMulti) return;
  const groups=lastMulti.groups, ctrl=groups.find(g=>/ck|对照|control|空白/i.test(g))||groups[groups.length-1];
  const del=gcDelSet('single');
  let rows=[['物质','中文名','类别'].concat(groups.flatMap(g=>[g+'-响应比',g+'-浓度(μg/L)',g+'-OAV'])).concat(['相对对照('+ctrl+'-rr比)','风味描述'])];
  const comps = (window.NatureViz && NatureViz.sortByCatValueDesc)
    ? NatureViz.sortByCatValueDesc(lastMulti.compounds, c=>maxRR(c))
    : [...lastMulti.compounds].sort((a,b)=>maxRR(b)-maxRR(a));
  comps.filter(c=>!del.has(c.en)).forEach(c=>{
    const name=dispName(c); let r=[c.en,name,c.cat||'其他'];
    groups.forEach(g=>{ const cell=c.by_group[g]; r.push(cell&&cell.rr!=null?cell.rr:''); r.push(cell&&cell._conc!=null?round2(cell._conc):''); r.push(cell&&cell._oav!=null?round2(cell._oav):''); });
    const cc=c.by_group[ctrl];
    if(cc&&cc.rr>0){ let parts=[]; groups.forEach(g=>{ if(g===ctrl)return; const cell=c.by_group[g]; if(cell&&cell.rr>0) parts.push(g+'='+(cell.rr/cc.rr).toFixed(2)); }); r.push(parts.join(';')); } else r.push('');
    r.push(c.odor||'');
    rows.push(r);
  });
  downloadCSV(rows,'处理组对比_'+groups.join('-')+'.csv');
}
function exportGroupCatCSV(){
  if(!lastMulti) return;
  const groups=lastMulti.groups, del=gcDelSet('single'); const catTotals={}; const catSet=[];
  groups.forEach(g=>catTotals[g]={});
  lastMulti.compounds.forEach(c=>{ if(del.has(c.en)) return; groups.forEach(g=>{ const cell=c.by_group[g]; if(cell&&cell._conc!=null){ const cat=c.cat||'其他'; catTotals[g][cat]=(catTotals[g][cat]||0)+cell._conc; if(catSet.indexOf(cat)<0)catSet.push(cat);} }); });
  catSet.sort((a,b)=>sumArr(catTotals,a)-sumArr(catTotals,b));
  let rows=[['类别'].concat(groups).concat(['合计'])];
  catSet.forEach(cat=>{ let row=[cat]; let t=0; groups.forEach(g=>{const v=catTotals[g][cat]||0;t+=v;row.push(round2(v));}); row.push(round2(t)); rows.push(row); });
  let tr=['合计']; let gt=0; groups.forEach(g=>{const v=catSet.reduce((s,cat)=>s+(catTotals[g][cat]||0),0);gt+=v;tr.push(round2(v));}); tr.push(round2(gt)); rows.push(tr);
  downloadCSV(rows,'处理组类别浓度总量.csv');
}
function round2(x){ return Math.round(x*100)/100; }
function downloadCSV(rows,fn){ let csv='\uFEFF'+rows.map(r=>r.map(c=>{ c=(c==null?'':c); c=String(c).replace(/"/g,'""'); return /[",\n]/.test(c)?'"'+c+'"':c; }).join(',')).join('\n');
  const blob=new Blob([csv],{type:'text/csv;charset=utf-8'}); const a=document.createElement('a'); a.href=URL.createObjectURL(blob); a.download=fn; a.click(); }

/* ===================== 多报告批量对比 ===================== */
let _samples=[];   // [{file:File, name:''}]
function onMultiFile(e){
  const f=e.target.files[0]; if(!f) return;
  _samples.push({file:f, name:''});
  renderSampleList();
  e.target.value='';  // 允许重复选同一文件
}
function renderSampleList(){
  const box=document.getElementById('sample-list');
  if(!_samples.length){box.innerHTML='<div class="field" style="flex-basis:100%"><span class="hint" style="color:var(--mut)">尚未添加报告，点击上方「选择文件」逐个添加。</span></div>';return;}
  box.innerHTML=_samples.map((s,i)=>`
    <div class="field" style="flex-basis:100%">
      <span>样品 ${i+1} 名称 <small style="color:var(--mut-2)">（${s.file.name}）</small></span>
      <div style="display:flex;gap:8px;align-items:center">
        <input type="text" placeholder="如 样品A / 发酵乳1号" value="${s.name}" oninput="_samples[${i}].name=this.value">
        <button class="ghost" style="flex:0 0 auto;padding:8px 12px" onclick="removeSample(${i})">删除</button>
      </div>
    </div>`).join('');
}
function removeSample(i){_samples.splice(i,1);renderSampleList();}
function clearMulti(){_samples=[];renderSampleList();document.getElementById('multi-out').innerHTML='';}

async function doMulti(){
  if(_samples.length<1){alert('请先添加至少 1 个报告文件');return;}
  const fd=new FormData();
  _samples.forEach(s=>fd.append('files', s.file));
  fd.append('ist_name',document.getElementById('m-ist').value.trim());
  _samples.forEach(s=>fd.append('names', s.name||''));
  document.getElementById('multi-out').innerHTML='<p class="hint">正在解析与对比 '+_samples.length+' 个报告…</p>';
  let data;
  try{ data=await(await fetch('/api/multi_upload',{method:'POST',body:fd})).json(); }
  catch(err){ document.getElementById('multi-out').innerHTML='<p class="hint">对比失败：'+err.message+'</p>'; return; }
  const samples=data.samples||[];
  if(!samples.length){document.getElementById('multi-out').innerHTML='<p class="hint">未解析到任何物质，请检查文件内容。</p>';return;}
  // 用后端自动识别的内标物质名预填（用户可手动改；手动改过则不再覆盖）
  const detectedIS=(samples.map(s=>s.ist_name).find(n=>n&&String(n).trim()))||'';
  if(detectedIS){ const mie=document.getElementById('m-ist'); if(mie && !mie.dataset.touched) mie.value=detectedIS; }
  // 统一内标参数（若填写）按公式 C=内标浓度×内标加量×响应比÷样品量 算浓度(μg/L)
  const cis=document.getElementById('m-cis').value.trim();
  const vis=document.getElementById('m-vis').value.trim();
  const ms=document.getElementById('m-ms').value.trim();
  const cisu=document.getElementById('m-cisu').value.trim()||'mg/L';
  const visu=document.getElementById('m-visu').value.trim()||'μL';
  const msu=document.getElementById('m-msu').value.trim()||'g';
  const ok = cis!==''&&vis!==''&&ms!==''&&!isNaN(parseFloat(cis))&&!isNaN(parseFloat(vis))&&!isNaN(parseFloat(ms))&&parseFloat(ms)!==0;
  samples.forEach(s=>{
    s.rows.forEach(r=>{
      if(ok){
        const resp=(typeof r.rr==='number'&&r.rr!=null)?r.rr:null;
        const c=resp!=null?conc(cis,vis,resp,ms,cisu,visu,msu):null;
        r.conc=c; r.conc_unit='μg/L'; r.conc_ugkg=c;
      } else if(typeof r.conc0==='number' && r.conc0!=null){ r.conc=r.conc0; r.conc_ugkg=r.conc0; r.conc_unit=(r.conc_unit||'μg/L'); }
      else { r.conc=null; r.conc_ugkg=null; }
    });
    computeOAV(s.rows);
    s.keyCount = s.rows.length;
    s.keyFrag = s.rows.filter(r=>r.oav_flag==='关键致香').length;
    s.maxOAV = s.rows.reduce((m,r)=>Math.max(m, (r.oav&&!isNaN(r.oav))?r.oav:0), 0);
  });
  renderMulti(samples);
}

let _multiSamples=null;  // 供导出/筛选复用
function renderMulti(samples){
  _multiSamples=samples;
  const fmtOAV=v=>(v==null||isNaN(v))?'—':v.toPrecision(3).replace(/\.?0+$/,'');
  // 1) 对比矩阵：以物质标准名聚合（已删除顶部「关键致香」KPI 模块）
  const map={};  // en -> {cn,cat,odor, bySample:{name:{conc,oav,roav,flag}}}
  samples.forEach(s=>{
    s.rows.forEach(r=>{
      const en=r.en; if(!en) return;
      if(!map[en]) map[en]={cn:r.cn, cat:r.cat, odor:r.odor, bySample:{}};
      map[en].bySample[s.name]={conc:r.conc, oav:r.oav, roav:r.roav, flag:r.oav_flag};
    });
  });
  const ensAll=Object.keys(map);
  // 类别集合（用于筛选下拉）+ 每类物质数
  const catMap={};
  ensAll.forEach(en=>{ const c=map[en].cat||'其他'; catMap[c]=(catMap[c]||0)+1; });
  const catOptions=Object.entries(catMap).sort((a,b)=>b[1]-a[1]);  // 按每类大小排列
  // 类别标准顺序（与数据库一致）：用 CAT 键序，未知类别排最后
  const catOrderKeys=Array.isArray(CAT)? CAT.map(c=>c.name||c) : Object.keys(CAT||{});
  const catRank=c=>{ const i=catOrderKeys.indexOf(c); return i>=0? i : 999; };
  // 单物质综合 OAV（取各样品最大值，用于同类内排序）
  const bestOAV=en=>Object.values(map[en].bySample).reduce((m,x)=>Math.max(m, (x&&x.oav!=null&&!isNaN(x.oav))?x.oav:0), 0);
  // 排序：先按类别顺序，同类内综合 OAV 由大到小
  const sortEns=(list)=>list.slice().sort((a,b)=>{
    const ca=map[a].cat||'其他', cb=map[b].cat||'其他';
    if(ca!==cb) return catRank(ca)-catRank(cb);
    return bestOAV(b)-bestOAV(a);
  });
  const nS=samples.length;
  const ctrlSample = samples.find(s=>/ck|对照|control|空白/i.test(s.name)) || samples[nS-1];
  const head='<th class="multi-del-col">✕</th><th>物质</th><th>中文名</th><th>类别</th>'+samples.map(s=>`<th>${escapeHtml(s.name)}<br><small>浓度 / OAV</small></th>`).join('')+`<th>相对于对照<br><small>${escapeHtml(ctrlSample?ctrlSample.name:'')}（浓度比）</small></th>`+'<th>气味活性</th><th>样品检出</th><th>风味描述</th>';
  // 渲染矩阵行的函数（按筛选类别）
  function buildRows(filterCat){
    const ens = filterCat? ensAll.filter(en=>(map[en].cat||'其他')===filterCat) : ensAll;
    const sorted = sortEns(ens);
    const del = mcDelSet();
    if(!sorted.length) return '<tr><td colspan="'+(nS+8)+'" style="text-align:center;color:var(--mut)">该类别下无物质</td></tr>';
    return sorted.map(en=>{
      if(del.has(en)) return '';                        // 已删除：直接从主表隐藏
      const m=map[en]; const cnt=Object.keys(m.bySample).length;
      // 中文列学习 GC-MS 解析逻辑：未收录（空 / (未收录)）时填入英文名，保证中文模块有内容
      const cnDisp = (m.cn && m.cn!=='(未收录)' && m.cn.trim()!=='')? m.cn : en;
      // 该物质整体气味活性（任一样品关键致香→关键；否则有潜在贡献→潜在；否则 —）
      let act='—', actCls='f-na';
      Object.values(m.bySample).forEach(x=>{ if(x&&x.flag==='关键致香'){act='关键致香';actCls='f-key';} });
      if(act==='—') Object.values(m.bySample).forEach(x=>{ if(x&&x.flag==='潜在贡献'){act='潜在贡献';actCls='f-pot';} });
      const cells=samples.map(s=>{
        const x=m.bySample[s.name];
        if(!x||x.conc==null||isNaN(x.conc)) return '<td class="m-none">—</td>';
        const cu=Number(x.conc).toPrecision(4).replace(/\.?0+$/,'');
        const fls=x.flag==='关键致香'?'f-key':(x.flag==='潜在贡献'?'f-pot':'f-na');
        return `<td class="oavflag ${fls} num-dark">${cu}<br><small>OAV ${fmtOAV(x.oav)}</small></td>`;
      }).join('');
      // 相对于对照：各样品浓度 ÷ 对照样品浓度（与单 PDF「相对对照」并列设计一致）
      let relCell='—';
      if(ctrlSample){
        const ccx=m.bySample[ctrlSample.name];
        if(ccx && typeof ccx.conc==='number' && ccx.conc>0){
          const parts=[];
          samples.forEach(s=>{ if(s===ctrlSample) return; const x=m.bySample[s.name];
            if(x && typeof x.conc==='number' && x.conc>0){ const fc=x.conc/ccx.conc;
              const cls=fc>=2?'fc-up':(fc<=0.5?'fc-down':''); const arr=fc>=2?'▲':(fc<=0.5?'▼':'');
              parts.push(`<span class="${cls}">${escapeHtml(s.name)} ${arr}${fc.toFixed(1)}×</span>`); } });
          relCell=parts.join('<br>')||'—';
        }
      }
      const od=(m.odor&&m.odor.trim())?m.odor:'—';
      return `<tr><td class="multi-del-col"><button class="multi-del-btn" data-en="${escapeHtml(en)}" title="删除该物质">✕</button></td><td><b>${en}</b></td><td class="cn">${cnDisp}</td><td>${badge(m.cat)}</td>${cells}<td style="text-align:left">${relCell}</td><td class="oavflag ${actCls}">${act}</td><td style="text-align:center">${cnt}/${nS}</td><td style="min-width:160px">${od}</td></tr>`;
    }).join('');
  }
  // 注：sortEns 返回新排序数组（不就地修改原数组）
  let filterBar='<div class="multi-toolbar">';
  filterBar+='<label>类别筛选：<select id="m-cat-filter" onchange="renderMultiMatrix()"><option value="">全部类别</option>';
  catOptions.forEach(([c,n])=>{ filterBar+=`<option value="${c}">${c}（${n}）</option>`; });
  filterBar+='</select></label>';
  filterBar+='<span class="hint">按类别顺序分组 · 同类数值由大到小</span>';
  filterBar+='<button class="ghost" onclick="exportMultiCSV()">⬇ 导出对比报告(CSV)</button>';
  filterBar+='<button class="ghost" onclick="exportMultiCatCSV()">⬇ 分类浓度总量(CSV)</button>';
  filterBar+='</div>';

  const tmp='<div id="multi-matrix-body"></div>';
  document.getElementById('multi-out').innerHTML =
    '<p class="hint">共解析 '+samples.length+' 个样品、'+ensAll.length+' 种风味物质参与对比（按类别分组、同类 OAV 由大到小排序，可点选类别筛选）。</p>' +
    filterBar + '<div class="wrap" id="multi-matrix"><table><thead><tr>'+head+'</tr></thead><tbody id="multi-matrix-rows">'+buildRows('')+'</tbody></table></div>' +
    renderMultiCatConc(samples) + '<div class="viz-wrap" id="viz-multi"></div>';
  // 暴露给筛选回调
  window.__buildMultiRows=buildRows;
  // 每行删除（小圆点 ✕）：与物质管理共享删除状态，删除后同步刷新主表与可视化
  const mm=document.getElementById('multi-matrix');
  if(mm && !mm._wired){ mm._wired=true; mm.addEventListener('click', function(e){ const b=e.target.closest('.multi-del-btn'); if(b) mcDeleteMulti(b.getAttribute('data-en')); }); }
  if(window.NatureViz) NatureViz.build('multi', _multiSamples);
}

// 各样品分类浓度总量对比：把每个样品内标法浓度(μg/L)按类别汇总，横向对比
function renderMultiCatConc(samples){
  if(!samples || !samples.length) return '';
  const perSample={};           // name -> {cat: totalConc}
  const catSet={};              // cat -> 全样品合计（用于排序）
  samples.forEach(s=>{
    const tot={};
    s.rows.forEach(r=>{
      const c=r.cat||'其他';
      const v=(typeof r.conc==='number'&&!isNaN(r.conc))?r.conc:0;
      tot[c]=(tot[c]||0)+v; catSet[c]=(catSet[c]||0)+v;
    });
    perSample[s.name]=tot;
  });
  const cats=Object.keys(catSet).sort((a,b)=>catSet[b]-catSet[a]);
  if(!cats.length) return '';
  const maxCat=Math.max.apply(null, cats.map(c=>Math.max.apply(null, samples.map(s=>(perSample[s.name][c]||0)))));
  let h='<div class="card" style="margin-top:16px"><h3><span class="ic">⚖</span>各样品分类浓度总量对比'+
        '<span class="sub">（内标法浓度 μg/L 按类别汇总）</span></h3><div class="scroll"><table class="multi-cat">'+
        '<thead><tr><th>类别</th>'+samples.map(s=>`<th>${escapeHtml(s.name)}</th>`).join('')+'<th>合计</th></tr></thead><tbody>';
  cats.forEach(c=>{
    let rowTot=0;
    const cells=samples.map(s=>{
      const v=perSample[s.name][c]||0; rowTot+=v;
      const w=maxCat? (v/maxCat*100):0;
      return `<td class="cc-cell"><span class="cc-bar" style="width:${w.toFixed(0)}%"></span><span class="cc-v">${v>0?v.toFixed(2):'—'}</span></td>`;
    }).join('');
    h+=`<tr><td>${badge(c)}</td>${cells}<td class="cc-sum">${rowTot.toFixed(2)}</td></tr>`;
  });
  const sumCells=samples.map(s=>{ let g=0; for(const c in perSample[s.name]) g+=perSample[s.name][c]; return `<td class="cc-sum">${g.toFixed(2)}</td>`; });
  let gAll=0; for(const c in catSet) gAll+=catSet[c];
  h+=`<tr class="cc-total"><td>合计</td>${sumCells}<td class="cc-sum">${gAll.toFixed(2)}</td></tr>`;
  h+='</tbody></table></div></div>';
  return h;
}
// 类别筛选时重绘矩阵行
function renderMultiMatrix(){
  const sel=document.getElementById('m-cat-filter');
  const f=sel?sel.value:'';
  const rows=window.__buildMultiRows?window.__buildMultiRows(f):'';
  const box=document.getElementById('multi-matrix-rows'); if(box) box.innerHTML=rows;
}
// 导出对比矩阵为 CSV
function exportMultiCSV(){
  const samples=_multiSamples||[]; if(!samples.length) return;
  const map={};
  samples.forEach(s=>s.rows.forEach(r=>{ if(!r.en) return; if(!map[r.en]) map[r.en]={cn:r.cn,cat:r.cat,odor:r.odor,bySample:{}}; map[r.en].bySample[s.name]={conc:r.conc,oav:r.oav,roav:r.roav}; }));
  const ens=Object.keys(map);
  const ctrlName=(samples.find(s=>/ck|对照|control|空白/i.test(s.name))||samples[samples.length-1]);
  const head=['物质','中文名','类别'].concat(samples.map(s=>s.name+'_浓度(μg/L)')).concat(samples.map(s=>s.name+'_OAV')).concat(samples.map(s=>s.name+'_ROAV')).concat(['相对于对照('+(ctrlName?ctrlName.name:'')+'_浓度比)','气味活性','样品检出数','风味描述']);
  const lines=[head.map(csvCell).join(',')];
  ens.forEach(en=>{
    const m=map[en]; const cnt=Object.keys(m.bySample).length;
    let act='—'; Object.values(m.bySample).forEach(x=>{ if(x&&x.flag==='关键致香')act='关键致香'; });
    if(act==='—') Object.values(m.bySample).forEach(x=>{ if(x&&x.flag==='潜在贡献')act='潜在贡献'; });
    const row=[en,m.cn||'',m.cat||'其他'];
    samples.forEach(s=>{ const x=m.bySample[s.name]; row.push((x&&x.conc!=null&&!isNaN(x.conc))?Number(x.conc).toPrecision(4).replace(/\.?0+$/,'') : '—'); });
    samples.forEach(s=>{ const x=m.bySample[s.name]; row.push((x&&x.oav!=null&&!isNaN(x.oav))?x.oav.toPrecision(3).replace(/\.?0+$/,'') : '—'); });
    samples.forEach(s=>{ const x=m.bySample[s.name]; row.push((x&&x.roav!=null&&!isNaN(x.roav))?x.roav.toFixed(2) : '—'); });
    let relParts=[]; if(ctrlName){ const ccx=m.bySample[ctrlName.name];
      if(ccx && typeof ccx.conc==='number' && ccx.conc>0){ samples.forEach(s=>{ if(s===ctrlName) return; const x=m.bySample[s.name];
        if(x && typeof x.conc==='number' && x.conc>0) relParts.push(s.name+'='+(x.conc/ccx.conc).toFixed(2)); }); } }
    row.push(relParts.join(';')||'');
    row.push(act);
    row.push(cnt+'/'+samples.length);
    row.push((m.odor&&m.odor.trim())?m.odor:'—');
    lines.push(row.map(csvCell).join(','));
  });
  const csv='\ufeff'+lines.join('\r\n');  // BOM 防中文乱码
  const url='data:text/csv;charset=utf-8,'+encodeURIComponent(csv);
  const a=document.createElement('a'); a.href=url; a.download='风味物质对比报告.csv'; a.click();
}
// 导出「各样品分类浓度总量」为 CSV（行=类别，列=样品 + 合计）
function exportMultiCatCSV(){
  const samples=_multiSamples||[]; if(!samples.length) return;
  const perSample={}; const catSet={};
  samples.forEach(s=>{
    const tot={};
    s.rows.forEach(r=>{ const c=r.cat||'其他'; const v=(typeof r.conc==='number'&&!isNaN(r.conc))?r.conc:0; tot[c]=(tot[c]||0)+v; catSet[c]=(catSet[c]||0)+v; });
    perSample[s.name]=tot;
  });
  const cats=Object.keys(catSet).sort((a,b)=>catSet[b]-catSet[a]);
  if(!cats.length) return;
  const head=['类别'].concat(samples.map(s=>s.name+'_浓度(μg/L)')).concat(['合计(μg/L)','占比(%)']);
  const lines=[head.map(csvCell).join(',')];
  cats.forEach(c=>{
    let rowTot=0; const row=[c];
    samples.forEach(s=>{ const v=perSample[s.name][c]||0; rowTot+=v; row.push(v>0?v.toFixed(2):'0'); });
    const gAll=Object.values(catSet).reduce((a,b)=>a+b,0);
    row.push(rowTot.toFixed(2));
    row.push(gAll? (rowTot/gAll*100).toFixed(1) : '0');
    lines.push(row.map(csvCell).join(','));
  });
  let gAll=0; for(const c in catSet) gAll+=catSet[c];
  lines.push(['合计'].concat(samples.map(s=>{ let g=0; for(const c in perSample[s.name]) g+=perSample[s.name][c]; return g.toFixed(2); })).concat([gAll.toFixed(2),'100.0']).map(csvCell).join(','));
  const csv='\ufeff'+lines.join('\r\n');
  const url='data:text/csv;charset=utf-8,'+encodeURIComponent(csv);
  const a=document.createElement('a'); a.href=url; a.download='分类浓度总量.csv'; a.click();
}
function csvCell(v){ v=(v==null)?'':String(v); if(/[",\r\n]/.test(v)) v='"'+v.replace(/"/g,'""')+'"'; return v; }
function escapeHtml(s){return (s||'').replace(/[&<>"]/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;'}[c]));}

let _allDB=null;
function _norm(s){return (s||'').toLowerCase().replace(/[\s\-(),.]/g,'');}
function populateBcat(){
  const sel=document.getElementById('bcat'); if(!sel||!_allDB) return;
  if(sel.dataset.filled) return;            // 仅首次填充，保留用户已选类别
  const cnt={}; _allDB.forEach(c=>{const k=c.cat||'其他'; cnt[k]=(cnt[k]||0)+1;});
  const opts=Object.entries(cnt).sort((a,b)=>b[1]-a[1]);
  sel.innerHTML='<option value="">全部类别</option>'+opts.map(([k,n])=>`<option value="${k}">${k}（${n}）</option>`).join('');
  sel.dataset.filled='1';
}
async function browseDB(){
  if(!_allDB){
    _allDB=await(await fetch('/api/all')).json();
  }
  populateBcat();
  const cat=document.getElementById('bcat').value;
  const key=_norm(document.getElementById('bkey').value);
  let rows=_allDB.filter(c=>(!cat||c.cat===cat) && (!key|| _norm(c.en).includes(key)||_norm(c.cn).includes(key)||_norm(c.odor).includes(key)||_norm(c.source).includes(key)|| (c.syn||[]).some(s=>_norm(s).includes(key))));
  // 信息完整度优先：有风味描述 + 阈值 + 来源 的物质排在前面
  rows=sortByCompleteness(rows);
  const nFull=rows.filter(r=>completeness(r)>=4).length;
  document.getElementById('bcount').textContent='共 '+rows.length+' 种（信息完整 '+nFull+' 种优先展示）';
  if(!rows.length){document.getElementById('browse-out').innerHTML='<p class="hint">无匹配</p>';return;}
  // 最多显示 300 条，避免过长（排序后前 300 条即信息最完整者）
  const show=rows.slice(0,300);
  let h='<div class="wrap" id="browse-wrap"><table><thead><tr><th class="c-idx">#</th><th class="c-en">英文名</th><th class="c-cn">中文名</th><th class="c-cas">CAS</th><th class="c-cat">类别</th><th class="c-thr">阈值(μg/L)</th><th class="c-med">介质</th><th class="c-match">匹配</th><th class="c-odor">气味描述</th><th class="c-src">来源</th></tr></thead><tbody>';
  show.forEach((r,i)=>{
    h+=`<tr><td class="c-idx">${i+1}</td><td class="c-en"><b>${r.en}</b></td><td class="c-cn cn">${r.cn??''}</td><td class="c-cas">${r.cas??''}</td><td class="c-cat">${badge(r.cat)}</td><td class="c-thr">${r.thr}</td><td class="c-med">${r.med}</td><td class="c-match">${r.match??''}</td><td class="c-odor">${r.odor??''}</td><td class="c-src src">${srcFmt(r.source)}</td></tr>`;
  });
  h+='</tbody></table></div>';
  if(rows.length>300) h+='<p class="hint">仅显示信息最完整的前 300 条，请缩小筛选条件查看全部。</p>';
  document.getElementById('browse-out').innerHTML=h;
}
// 初始自动加载数据库列表（无需手动输入即见数据）
setTimeout(browseDB, 300);

function exportCSV(which){
  const rows=which==='report'?lastReport:lastBatch;if(!rows.length){alert('请先检索');return;}
  const head=['英文名','中文名','CAS','类别','RT','含量(响应比)','阈值(μg/L)','介质','浓度(μg/L)','OAV','ROAV','气味活性','匹配方式','气味描述','来源'];
  const lines=[head.join(',')];
  rows.forEach(r=>{
    const cu = (r.conc==null||isNaN(r.conc))?'':r.conc;
    const oav = (r.oav==null)?'':r.oav;
    const roav = (r.roav==null)?'':r.roav;
    lines.push([r.en,r.cn,r.cas??'',r.cat,r.rt??'',r.rr??'',r.thr,r.med??'',cu,oav,roav,r.oav_flag??'',r.match,r.odor??'',String(r.source??'').replace(/\n/g,'、')].map(x=>`"${String(x).replace(/"/g,'""')}"`).join(','));
  });
  const blob=new Blob(['\ufeff'+lines.join('\n')],{type:'text/csv;charset=utf-8'});
  const a=document.createElement('a');a.href=URL.createObjectURL(blob);a.download='风味物质检索结果.csv';a.click();
}

fetch('/api/info').then(r=>r.json()).then(d=>document.getElementById('dbsize').textContent=d.size);
// 调试用：URL 加 #demo-batch 自动填批量框并触发检索
if(location.hash.startsWith('#demo-batch')){
  setTimeout(function(){
    var tab=document.querySelector('.tab[data-t="batch"]'); if(tab) tab.click();
    var ta=document.getElementById('batch');
    ta.value=['Butanoic acid','Hexanal','2-Heptanone','gamma-Nonalactone','Limonene','Linalool','Ethyl butanoate','Furfurylthiol','Acetic acid','2-Phenylethanol','Benzaldehyde','Furaneol'].join('\n');
    if(typeof doBatch==='function') doBatch();
  }, 400);
}
</script>
<script>
/* ===== NATURE_VIZ_START ===== */
/* Nature 风格可视化（热力图 / PCA）+ 物质管理 */
(function(){
  'use strict';
  var NV = {};
  NV.palette = ['#24908c','#e0844a','#4a7de0','#b0556f','#7a9b3b','#9b6bb0','#d09a2e','#3aa0a0','#c05a5a','#5a7fc0','#6fae5a','#c98a3a'];
  // 类别配色：沿用 CAT 顺序，未命中则用名称哈希兜底
  NV.catColor = function(cat){
    try{
      var order = (typeof CAT!=='undefined' && CAT) ? (Array.isArray(CAT)?CAT.map(function(c){return c.name||c;}):Object.keys(CAT)) : [];
      var i = order.indexOf(cat); if(i<0){ var s=0; for(var k=0;k<(''+cat).length;k++) s=(s*31+cat.charCodeAt(k))>>>0; i=s; }
      return NV.palette[i % NV.palette.length];
    }catch(e){ return NV.palette[0]; }
  };
  // viridis 近似色图（向后兼容保留）
  var VSTOPS=[[68,1,84],[59,82,139],[33,145,140],[94,201,98],[253,231,37]];
  NV.viridis = function(t){
    t = Math.max(0,Math.min(1,t)); var n=VSTOPS.length-1, f=t*n, i=Math.floor(f), fr=f-i;
    if(i>=n) i=n-1, fr=1; var a=VSTOPS[i], b=VSTOPS[i+1];
    return 'rgb('+Math.round(a[0]+(b[0]-a[0])*fr)+','+Math.round(a[1]+(b[1]-a[1])*fr)+','+Math.round(a[2]+(b[2]-a[2])*fr)+')';
  };
  // RdBu_r 红-白-蓝 发散色图（Nature 同款），输入 t∈[-1,1]；0=白，1=红，-1=蓝
  var RBSTOPS=[[33,102,172],[103,169,207],[209,229,240],[253,219,199],[239,138,98],[178,24,43]];
  NV.rdbu = function(t){
    t = Math.max(-1,Math.min(1,t)); var tt=(t+1)/2, n=RBSTOPS.length-1, f=tt*n, i=Math.floor(f), fr=f-i;
    if(i>=n) i=n-1, fr=1; var a=RBSTOPS[i], b=RBSTOPS[i+1];
    return 'rgb('+Math.round(a[0]+(b[0]-a[0])*fr)+','+Math.round(a[1]+(b[1]-a[1])*fr)+','+Math.round(a[2]+(b[2]-a[2])*fr)+')';
  };
  // —— 热力图可选配色方案（每组为按 t∈[0,1] 排列的 RGB 端点；0=低值, 1=高值）—— 默认蓝橙（#88C2EA→#FFA74F）
  var NV_PALETTES={
    'blueorange':{name:'蓝橙',   stops:[[136,194,234],[190,216,240],[247,243,236],[255,201,150],[255,167,79]]},
    'rdbu':      {name:'红蓝',   stops:[[33,102,172],[103,169,207],[247,247,247],[239,138,98],[178,24,43]]},
    'bluepink':  {name:'蓝粉',   stops:[[220,239,251],[120,180,235],[237,233,245],[244,180,205],[224,81,142]]},
    'pureblue':  {name:'纯蓝',   stops:[[240,248,255],[173,216,255],[96,160,240],[36,96,200],[12,46,120]]},
    'gray':      {name:'灰阶',   stops:[[247,247,247],[200,200,200],[120,120,120],[55,55,55]]}
  };
  NV._palette=(function(){ try{ var p=localStorage.getItem('gcms_nv_palette')||'blueorange'; return NV_PALETTES[p]?p:'blueorange'; }catch(e){ return 'blueorange'; } })();
  NV.paletteStops=function(){ var p=NV_PALETTES[NV._palette]||NV_PALETTES.blueorange; return p.stops; };
  NV.bp = function(t){
    var ST=NV.paletteStops();
    t = Math.max(0,Math.min(1,t)); var n=ST.length-1, f=t*n, i=Math.floor(f), fr=f-i;
    if(i>=n) i=n-1, fr=1; var a=ST[i], b=ST[i+1];
    return 'rgb('+Math.round(a[0]+(b[0]-a[0])*fr)+','+Math.round(a[1]+(b[1]-a[1])*fr)+','+Math.round(a[2]+(b[2]-a[2])*fr)+')';
  };
  NV.fmt = function(v){ if(v==null||!isFinite(v)) return '—'; var a=Math.abs(v);
    if(a>=1e6) return (v/1e6).toFixed(2)+'M'; if(a>=1e3) return (v/1e3).toFixed(1)+'k';
    if(a>=100) return ''+Math.round(v); return (''+Number(v).toPrecision(3)).replace(/\.?0+$/,''); };

  /* —— 记忆版排序：按类别分组（沿用 CAT 顺序），组内按数值降序 ——
     供「文件解析 → 分类物质」等后续功能直接复用：NV.sortByCatValueDesc(items, valFn) */
  NV.sortByCatValueDesc = function(items, valFn){
    var order = (typeof CAT!=='undefined' && CAT) ? (Array.isArray(CAT)?CAT.map(function(c){return c.name||c;}):Object.keys(CAT)) : [];
    var rank = function(c){ var i=order.indexOf(c); return i>=0?i:999; };
    var groups = {};
    items.forEach(function(it){ var c=it.cat||'其他'; (groups[c]=groups[c]||[]).push(it); });
    var cats = Object.keys(groups).sort(function(a,b){ return rank(a)-rank(b); });
    var out=[];
    cats.forEach(function(c){ groups[c].sort(function(x,y){ return (valFn(y)||0)-(valFn(x)||0); }); out=out.concat(groups[c]); });
    return out;
  };

  /* —— PCA（取前 2 主成分，幂迭代 + 收缩） —— */
  function center(M){ var n=M.length,p=M[0].length,mn=M[0].map(function(_,j){return M.reduce(function(s,r){return s+r[j];},0)/n;});
    return M.map(function(r){return r.map(function(v,j){return v-mn[j];});}); }
  function cov(Mc){ var n=Mc.length,p=Mc[0].length,C=[]; for(var a=0;a<p;a++){C.push(new Array(p).fill(0));}
    for(var i=0;i<n;i++) for(var a=0;a<p;a++) for(var b=0;b<p;b++) C[a][b]+=Mc[i][a]*Mc[i][b];
    for(var a=0;a<p;a++) for(var b=0;b<p;b++) C[a][b]/=(n-1); return C; }
  function mv(C,v){ return C.map(function(row){return row.reduce(function(s,x,i){return s+x*v[i];},0);}); }
  function power(C,v0,it){ var v=v0.slice(); for(var k=0;k<it;k++){ var nw=mv(C,v),norm=Math.hypot.apply(null,nw)||1; v=nw.map(function(x){return x/norm;}); } return v; }
  NV.pca2 = function(M){
    var n=M.length, p=M[0]?M[0].length:0;
    if(n<1||p<1) return {scores:M.map(function(){return [0,0];}),evr:[0,0],v1:[1,0],v2:[0,1]};
    if(p===1){ var Mc=center(M); return {scores:Mc.map(function(r){return [r[0],0];}),evr:[1,0],v1:[1],v2:[0]}; }
    var Mc=center(M), C=cov(Mc);
    var v1=power(C, M[0].map(function(_,i){return Math.cos(i+1);}),80);
    var l1=mv(C,v1).reduce(function(s,x,i){return s+x*v1[i];},0);
    var C2=C.map(function(row,a){return row.map(function(x,b){return x-l1*v1[a]*v1[b];});});
    var v2=power(C2, M[0].map(function(_,i){return Math.sin(i+2);}),80);
    var l2=mv(C2,v2).reduce(function(s,x,i){return s+x*v2[i];},0);
    var scores=Mc.map(function(r){return [r.reduce(function(s,x,i){return s+x*v1[i];},0), r.reduce(function(s,x,i){return s+x*v2[i];},0)];});
    var tot=0; for(var a=0;a<p;a++) tot+=C[a][a];
    return {scores:scores, evr:[l1/(tot||1), l2/(tot||1)], v1:v1, v2:v2};
  };

  /* —— 数据归一化 —— */
  NV.singleItems = function(ds, metric){
    return (ds.compounds||[]).map(function(c){
      var vals={}; (ds.groups||[]).forEach(function(g){ var cell=c.by_group[g]; vals[g]=(cell!=null)?(metric==='conc'?cell._conc:cell._oav):undefined; });
      var mv=0; (ds.groups||[]).forEach(function(g){ if(vals[g]!=null&&isFinite(vals[g])) mv=Math.max(mv,vals[g]); });
      return {en:c.en, cn:(typeof dispName==='function'?dispName(c):(c.cn||c.en)), cat:c.cat||'其他', odor:c.odor||'', vals:vals, maxVal:mv};
    });
  };
  NV.multiItems = function(samples, metric){
    var map={};
    (samples||[]).forEach(function(s){ (s.rows||[]).forEach(function(r){ if(!r.en) return;
      if(!map[r.en]) map[r.en]={en:r.en,cn:r.cn,cat:r.cat||'其他',odor:r.odor||'',vals:{},maxVal:0};
      // 跨样品择优：若已存为「未收录」而本样品有匹配中文名，则采用更优的中文名
      else if((!map[r.en].cn || map[r.en].cn==='(未收录)') && r.cn && r.cn!=='(未收录)' && r.cn.trim()!=='') map[r.en].cn=r.cn;
      map[r.en].vals[s.name]=(metric==='conc'?r.conc:r.oav); }); });
    return Object.keys(map).map(function(k){ var m=map[k]; var mv=0; for(var g in m.vals){ if(m.vals[g]!=null&&isFinite(m.vals[g])) mv=Math.max(mv,m.vals[g]); } m.maxVal=mv;
      // 未收录（空 / (未收录)）时退回英文名，与解析主表（中文列学习 GC-MS 逻辑）保持一致
      if(!m.cn || m.cn==='(未收录)' || m.cn.trim()==='') m.cn=m.en; return m; });
  };

  /* —— 热力图（SVG：先按行 min-max 归一化，低饱和蓝→粉 色系，Q 版圆角方块 cell） —— */
  NV.heatmap = function(host, o){
    var rows=o.rows, cols=o.cols;
    if(!rows||!rows.length||!cols||!cols.length){ host.innerHTML='<p class="hint">无可用数据（请检查类别 / 组别选择与删除设置）</p>'; return; }
    // 预处理 + 每行归一化到 [0,1]（使各物质跨组相对模式可比）
    var data=rows.map(function(it){
      var raw=cols.map(function(g){ var v=it.vals[g]; return (v!=null&&isFinite(v))?v:null; });
      var present=raw.filter(function(v){return v!=null;});
      var rmin=present.length?Math.min.apply(null,present):0, rmax=present.length?Math.max.apply(null,present):1, rspan=(rmax-rmin)||1;
      var nz=raw.map(function(v){ return (v==null)?null:(v-rmin)/rspan; });
      return {it:it, raw:raw, nz:nz};
    });
    var all=[]; data.forEach(function(d){ d.raw.forEach(function(v){ if(v!=null) all.push(v); }); });
    if(!all.length){ host.innerHTML='<p class="hint">所选物质在所选组别下无数值</p>'; return; }
    var mL=158, mT=70, mR=70, mB=18;
    var cW=Math.max(48, Math.min(110, Math.round(420/cols.length))), cH=Math.max(18, Math.min(28, Math.round(620/rows.length)));
    var W=mL+cols.length*cW+mR, H=mT+rows.length*cH+mB;
    var s='<svg viewBox="0 0 '+W+' '+H+'" role="img">';
    // 色条（蓝色系活力分色，offset 0=底=低，1=顶=高）
    var cbX=W-mR+18, cbY=mT, cbH=Math.min(160, H-mT-mB), cbW=12;
    s+='<defs><linearGradient id="nvcb" x1="0" y1="1" x2="0" y2="0">';
    for(var t=0;t<=10;t++){ var tt=t/10; s+='<stop offset="'+(t/10)+'" stop-color="'+NV.bp(tt)+'"/>'; }
    s+='</linearGradient></defs>';
    s+='<rect x="'+cbX+'" y="'+cbY+'" width="'+cbW+'" height="'+cbH+'" rx="3" fill="url(#nvcb)" stroke="#d7dde6"/>';
    s+='<text class="viz-cbar-t" x="'+(cbX+cbW+5)+'" y="'+(cbY+9)+'">高</text>';
    s+='<text class="viz-cbar-t" x="'+(cbX+cbW+5)+'" y="'+(cbY+cbH/2+3)+'">中</text>';
    s+='<text class="viz-cbar-t" x="'+(cbX+cbW+5)+'" y="'+(cbY+cbH)+'">低</text>';
    // 列标签（旋转）
    cols.forEach(function(g,j){ var x=mL+j*cW+cW/2;
      s+='<text class="viz-clabel" transform="translate('+(x+4)+','+(mT-12)+') rotate(-40)" text-anchor="start">'+NV.esc(g)+'</text>'; });
    // 行（Q 版圆角方块）
    var pad=Math.max(1.5, Math.min(cW,cH)*0.10), rx=Math.min(cW,cH)*0.32;
    data.forEach(function(d,i){ var it=d.it, y=mT+i*cH;
      s+='<rect x="6" y="'+(y+2)+'" width="6" height="'+(cH-4)+'" fill="'+NV.catColor(it.cat)+'"/>';
      s+='<text class="viz-rlabel" x="'+(mL-8)+'" y="'+(y+cH/2+4)+'" text-anchor="end">'+NV.esc(it.cn||it.en)+'</text>';
      d.nz.forEach(function(z,j){ var cx=mL+j*cW, cy=y;
        if(z==null){ s+='<rect x="'+(cx+pad)+'" y="'+(cy+pad)+'" width="'+(cW-2*pad)+'" height="'+(cH-2*pad)+'" rx="'+rx+'" fill="#f1f3f7" stroke="#e3e8ef" stroke-width="0.5"/>'; }
        else { s+='<rect x="'+(cx+pad)+'" y="'+(cy+pad)+'" width="'+(cW-2*pad)+'" height="'+(cH-2*pad)+'" rx="'+rx+'" fill="'+NV.bp(z)+'" stroke="#ffffff" stroke-width="0.6"/>'; }
      });
    });
    s+='</svg>';
    host.innerHTML=s;
  };

  /* —— 热力图导出：SVG（矢量）/ PNG（2× 栅格，白底，内联文字样式） —— */
  NV._dl=function(blob, name){ try{ var url=URL.createObjectURL(blob); var a=document.createElement('a'); a.href=url; a.download=name; document.body.appendChild(a); a.click(); setTimeout(function(){ try{document.body.removeChild(a);}catch(_){} URL.revokeObjectURL(url); }, 120); }catch(e){ alert('导出失败：'+e.message); } };
  NV.exportFig=function(kind, fmt){
    var host=document.getElementById('viz-hm-'+kind); if(!host) return;
    var svg=host.querySelector('svg'); if(!svg){ alert('暂无可导出的热力图，请先完成解析'); return; }
    var vb=(svg.getAttribute('viewBox')||'0 0 600 400').split(/\s+/);
    var W=parseFloat(vb[2])||600, H=parseFloat(vb[3])||400;
    var clone=svg.cloneNode(true);
    clone.setAttribute('xmlns','http://www.w3.org/2000/svg');
    clone.setAttribute('width', W); clone.setAttribute('height', H);
    // 内联关键文字样式（脱离页面 CSS 后仍能正确显示）
    var style=document.createElementNS('http://www.w3.org/2000/svg','style');
    style.textContent='text{font-family:Arial,"PingFang SC","Microsoft YaHei",sans-serif}'
      +'.viz-rlabel{fill:#2a3340;font-size:11px}.viz-clabel{fill:#5b6776;font-size:11px;font-weight:600}'
      +'.viz-cbar-t{fill:#5b6776;font-size:10.5px}.viz-axis-t{fill:#5b6776;font-size:11px}';
    clone.insertBefore(style, clone.firstChild);
    // 白底
    var bg=document.createElementNS('http://www.w3.org/2000/svg','rect');
    bg.setAttribute('x',0); bg.setAttribute('y',0); bg.setAttribute('width',W); bg.setAttribute('height',H); bg.setAttribute('fill','#ffffff');
    clone.insertBefore(bg, clone.firstChild);
    var xml=new XMLSerializer().serializeToString(clone);
    if(fmt==='svg'){ NV._dl(new Blob([xml],{type:'image/svg+xml;charset=utf-8'}), 'heatmap_'+kind+'.svg'); return; }
    var img=new Image();
    img.onload=function(){
      var scale=2, cv=document.createElement('canvas'); cv.width=W*scale; cv.height=H*scale;
      var cx=cv.getContext('2d'); cx.fillStyle='#ffffff'; cx.fillRect(0,0,cv.width,cv.height);
      cx.drawImage(img,0,0,cv.width,cv.height);
      cv.toBlob(function(b){ if(b) NV._dl(b,'heatmap_'+kind+'.png'); else alert('PNG 导出失败，可改用 SVG'); },'image/png');
    };
    img.onerror=function(){ alert('PNG 导出失败，可改用 SVG 导出'); };
    img.src='data:image/svg+xml;charset=utf-8,'+encodeURIComponent(xml);
  };

  /* —— PCA 得分图（SVG）：按组别着色 + 置信椭圆 + rug 边缘须图 —— */
  NV.groupColor = function(g, idx){
    // 组别配色：用 NV.palette 前若干色循环；同一组名同一色
    var key=''+(g==null?'__':g);
    if(!NV._groupColorCache) NV._groupColorCache={};
    if(NV._groupColorCache[key]) return NV._groupColorCache[key];
    var c=NV.palette[(idx||Object.keys(NV._groupColorCache).length) % NV.palette.length];
    NV._groupColorCache[key]=c; return c;
  };
  NV.resetGroupColorCache = function(){ NV._groupColorCache={}; };
  NV.pcaPlot = function(host, o){
    var pts=o.points||[];
    if(!pts.length){ host.innerHTML='<p class="hint">无可用数据</p>'; return; }
    var xs=pts.map(function(p){return p.x;}), ys=pts.map(function(p){return p.y;});
    var xmn=Math.min.apply(null,xs), xmx=Math.max.apply(null,xs), ymn=Math.min.apply(null,ys), ymx=Math.max.apply(null,ys);
    var xd=(xmx-xmn)||1, yd=(ymx-ymn)||1, pad=Math.max(xd,yd)*0.18+0.01;
    xmn-=pad; xmx+=pad; ymn-=pad; ymx+=pad;
    var mL=58, mR=30, mT=24, mB=52;
    var W=Math.max(440,560), H=Math.max(360,440);
    var sx=function(x){ return mL+(x-xmn)/(xmx-xmn)*(W-mL-mR); };
    var sy=function(y){ return mT+(1-(y-ymn)/(ymx-ymn))*(H-mT-mB); };
    var evr=o.evr||[0,0];
    var s='<svg viewBox="0 0 '+W+' '+H+'" role="img">';
    // 0 轴虚线（仿 Nature 风格）
    function dashline(x1,y1,x2,y2){
      s+='<line x1="'+x1+'" y1="'+y1+'" x2="'+x2+'" y2="'+y2+'" stroke="#9aa6b4" stroke-width="0.7" stroke-dasharray="3 3"/>';
    }
    if(xmn<0 && xmx>0) dashline(sx(0),sy(ymn),sx(0),sy(ymx));
    if(ymn<0 && ymx>0) dashline(sx(xmn),sy(0),sx(xmx),sy(0));
    // 轴框
    s+='<line x1="'+mL+'" y1="'+sy(ymn)+'" x2="'+mL+'" y2="'+sy(ymx)+'" stroke="#3a4a5c" stroke-width="0.9"/>';
    s+='<line x1="'+mL+'" y1="'+(H-mB)+'" x2="'+sx(xmx)+'" y2="'+(H-mB)+'" stroke="#3a4a5c" stroke-width="0.9"/>';
    // rug（每个点一条小竖线）
    pts.forEach(function(p){ var c=NV.groupColor(p.grp, p.gi||0);
      s+='<line x1="'+sx(p.x)+'" y1="'+sy(ymn)+'" x2="'+sx(p.x)+'" y2="'+(sy(ymn)+Math.max(4,(H-mT-mB)*0.025))+'" stroke="'+c+'" stroke-width="0.6" opacity="0.55"/>';
      s+='<line x1="'+mL+'" y1="'+sy(p.y)+'" x2="'+(mL-Math.max(4,(W-mL-mR)*0.025))+'" y2="'+sy(p.y)+'" stroke="'+c+'" stroke-width="0.6" opacity="0.55"/>';
    });
    // 置信椭圆（按组）
    var groups={};
    pts.forEach(function(p){ var g=p.grp||'—'; (groups[g]=groups[g]||[]).push(p); });
    var grpOrder=Object.keys(groups);
    Object.keys(groups).forEach(function(g,gi){
      var arr=groups[g]; if(arr.length<2) return;
      var mx=arr.reduce(function(s,p){return s+p.x;},0)/arr.length;
      var my=arr.reduce(function(s,p){return s+p.y;},0)/arr.length;
      var sxx=arr.reduce(function(s,p){return s+(p.x-mx)*(p.x-mx);},0)/arr.length;
      var syy=arr.reduce(function(s,p){return s+(p.y-my)*(p.y-my);},0)/arr.length;
      var sxy=arr.reduce(function(s,p){return s+(p.x-mx)*(p.y-my);},0)/arr.length;
      var det=sxx*syy-sxy*sxy; if(det<=0) return;
      var a=Math.sqrt(Math.max(0, (sxx+syy+Math.sqrt((sxx-syy)*(sxx-syy)+4*sxy*sxy))/2));
      var b=Math.sqrt(Math.max(0, (sxx+syy-Math.sqrt((sxx-syy)*(sxx-syy)+4*sxy*sxy))/2));
      var theta=Math.atan2(2*sxy, sxx-syy)/2;
      // 95% 置信椭圆：k²=5.991（χ²(2)），等效放大因子 sqrt(5.991)*σ
      var k=Math.sqrt(5.991);
      var c=NV.groupColor(g, gi);
      var path=''; var n=48;
      for(var i=0;i<=n;i++){
        var ang=2*Math.PI*i/n;
        var ex=a*k*Math.cos(ang), ey=b*k*Math.sin(ang);
        var xr=mx + ex*Math.cos(theta) - ey*Math.sin(theta);
        var yr=my + ex*Math.sin(theta) + ey*Math.cos(theta);
        path+=(i===0?'M':'L')+sx(xr)+' '+sy(yr)+' ';
      }
      s+='<path d="'+path+'Z" fill="'+c+'" fill-opacity="0.10" stroke="'+c+'" stroke-width="1.2" stroke-opacity="0.85"/>';
    });
    // 散点
    pts.forEach(function(p){
      var c=NV.groupColor(p.grp, p.gi||0);
      var shape=p.shape||'circle';
      if(shape==='square') s+='<rect x="'+(sx(p.x)-3.2)+'" y="'+(sy(p.y)-3.2)+'" width="6.4" height="6.4" fill="'+c+'" fill-opacity="0.88" stroke="#fff" stroke-width="0.7"/>';
      else if(shape==='triangle') s+='<polygon points="'+sx(p.x)+','+(sy(p.y)-4)+' '+(sx(p.x)-4)+','+(sy(p.y)+3.2)+' '+(sx(p.x)+4)+','+(sy(p.y)+3.2)+'" fill="'+c+'" fill-opacity="0.88" stroke="#fff" stroke-width="0.7"/>';
      else s+='<circle cx="'+sx(p.x)+'" cy="'+sy(p.y)+'" r="3.6" fill="'+c+'" fill-opacity="0.9" stroke="#fff" stroke-width="0.7"/>';
    });
    if(pts.length<=30){ pts.forEach(function(p){ s+='<text x="'+(sx(p.x)+5)+'" y="'+(sy(p.y)+3)+'" font-size="9" fill="#23303c">'+NV.esc(p.label||'')+'</text>'; }); }
    // 轴标题（带方差贡献率，仿 Nature）
    s+='<text class="viz-axis-t" x="'+(mL+(W-mL-mR)/2)+'" y="'+(H-10)+'" text-anchor="middle" font-size="12" fill="#1f2a3d" font-weight="700">PC1 ('+(evr[0]*100).toFixed(1)+'%)</text>';
    s+='<text class="viz-axis-t" transform="translate('+(mL-32)+','+(mT+(H-mT-mB)/2)+') rotate(-90)" text-anchor="middle" font-size="12" fill="#1f2a3d" font-weight="700">PC2 ('+(evr[1]*100).toFixed(1)+'%)</text>';
    // 刻度
    var tk=4;
    var xstep=(xmx-xmn)/tk, ystep=(ymx-ymn)/tk;
    for(var i=1;i<tk;i++){
      var xv=xmn+i*xstep, yv=ymn+i*ystep;
      s+='<text x="'+sx(xv)+'" y="'+(H-mB+12)+'" font-size="9" fill="#5d6b7c" text-anchor="middle">'+NV.fmt(xv)+'</text>';
      s+='<text x="'+(mL-6)+'" y="'+(sy(yv)+3)+'" font-size="9" fill="#5d6b7c" text-anchor="end">'+NV.fmt(yv)+'</text>';
    }
    s+='</svg>';
    host.innerHTML=s;
    // 图例（按组别）
    var leg=document.getElementById(o.legendId);
    if(leg){
      var h='<b class="leg-title">Groups</b> ';
      grpOrder.forEach(function(g,gi){ var c=NV.groupColor(g,gi);
        h+='<span><i style="background:'+c+'"></i>'+NV.esc(g)+' ('+groups[g].length+')</span>'; });
      leg.innerHTML=h;
    }
  };

  NV.esc = function(s){ return (s==null?'':(''+s)).replace(/[&<>"]/g,function(c){return {'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;'}[c];}); };

  /* —— 状态持久化 —— */
  NV.KEY=function(kind){ return 'gcms_nv_'+kind; };
  NV.load=function(kind){ try{ var raw=localStorage.getItem(NV.KEY(kind)); return raw?JSON.parse(raw):null; }catch(e){ return null; } };
  NV.save=function(kind){ try{ var st=NV.state[kind]; if(!st) return; localStorage.setItem(NV.KEY(kind), JSON.stringify({
    deleted:[].slice.call(st.deleted), added:[].slice.call(st.added), addedItems:st.addedItems||{}, metric:st.metric,
    cats:[].slice.call(st.cats), groups:[].slice.call(st.groups) })); }catch(e){} };

  NV.template=function(kind){ return ''
    +'<div class="viz-head"><h3><span class="ic">📊</span>可视化分析 · 热力图</h3>'
    +'<span class="hint">勾选/删除物质、切换类别与组别、切换指标，右侧物质管理实时联动热力图</span></div>'
    +'<div class="viz-toolbar">'
    +'<label>指标</label><select id="viz-metric-'+kind+'"><option value="oav">OAV</option><option value="conc">浓度(μg/L)</option></select>'
    +'<label>配色</label><select id="viz-palette-'+kind+'">'+Object.keys(NV_PALETTES).map(function(k){return '<option value="'+k+'"'+(k===NV._palette?' selected':'')+'>'+NV_PALETTES[k].name+'</option>';}).join('')+'</select>'
    +'<input class="viz-search" id="viz-search-'+kind+'" placeholder="搜索物质名…">'
    +'<button class="ghost" id="viz-add-'+kind+'">＋库内新增</button>'
    +'<button class="ghost" id="viz-save-'+kind+'">保存调整</button>'
    +'<button class="ghost" id="viz-undo-'+kind+'">撤回删除</button>'
    +'<button class="ghost" id="viz-reset-'+kind+'">重置</button>'
    +'<span class="viz-sep"></span>'
    +'<button class="ghost" id="viz-png-'+kind+'">导出 PNG</button>'
    +'<button class="ghost" id="viz-svg-'+kind+'">导出 SVG</button></div>'
    +'<div class="viz-chips" id="viz-cat-'+kind+'"></div>'
    +'<div class="viz-chips" id="viz-grp-'+kind+'"></div>'
    +'<div class="viz-main">'
    +'<div class="viz-fig"><h4>热力图 <small id="viz-hm-sub-'+kind+'"></small></h4><div id="viz-hm-'+kind+'"></div><div class="viz-legend" id="viz-hm-leg-'+kind+'"></div></div>'
    +'<div class="viz-mgr"><div class="mgr-top"><b>物质管理</b>'
    +'<button class="ghost" id="viz-all-'+kind+'">全选</button><button class="ghost" id="viz-none-'+kind+'">全不选</button>'
    +'<span class="hint" id="viz-mgr-stat-'+kind+'"></span></div>'
    +'<div class="mgr-list" id="viz-mgr-'+kind+'"></div>'
    +'<div class="viz-hint">勾选=纳入分析；取消=删除（可「撤回删除」）。「保存调整」将删除设置写入浏览器本地，下次同类型解析保留。可用「＋库内新增」按名称加入库内物质。</div></div>'
    +'</div>'; };

  NV.build=function(kind, ds){
    var host=document.getElementById('viz-'+kind); if(!host) return;
    try{
      var groupsAvail = (kind==='single') ? (ds.groups||[]).slice() : (ds||[]).map(function(s){return s.name;});
      var stored=NV.load(kind);
      var itemsAll = (kind==='single')? NV.singleItems(ds, (stored&&stored.metric)||'oav') : NV.multiItems(ds, (stored&&stored.metric)||'oav');
      var catsAvail=[].concat.apply([],[]); catsAvail=[...new Set(itemsAll.map(function(it){return it.cat||'其他';}))];
      var st={ metric:(stored&&stored.metric)||'oav', cats:new Set(), groups:new Set(groupsAvail),
        deleted:new Set(stored&&stored.deleted?stored.deleted:[]), added:new Set(stored&&stored.added?stored.added:[]),
        addedItems: (stored&&stored.addedItems)||{}, undoStack:[], search:'', _ds:ds, _groupsAvail:groupsAvail, mgrCollapsed:new Set(catsAvail) };
      st.cats = new Set((stored&&stored.cats&&stored.cats.length)? stored.cats.filter(function(c){return catsAvail.indexOf(c)>=0;}) : catsAvail);
      NV.state[kind]=st;
      host.innerHTML=NV.template(kind);
      NV.wire(kind);
      NV.render(kind);
    }catch(e){ console.error('NatureViz build error',e); try{host.innerHTML='<p class="hint">可视化渲染失败：'+e.message+'</p>';}catch(_){} }
  };

  NV.render=function(kind){
    var st=NV.state[kind]; if(!st) return; var ds=st._ds, metric=st.metric;
    var itemsAll = (kind==='single')? NV.singleItems(ds, metric) : NV.multiItems(ds, metric);
    var selGroups = st._groupsAvail.filter(function(g){return st.groups.has(g);});
    var q=st.search.trim().toLowerCase();
    var cand=itemsAll.filter(function(it){
      return st.cats.has(it.cat) && selGroups.some(function(g){return it.vals[g]!=null&&isFinite(it.vals[g]);})
        && (!q || (it.cn&&it.cn.toLowerCase().indexOf(q)>=0) || (it.en&&it.en.toLowerCase().indexOf(q)>=0)); });
    var enInc=new Set(cand.map(function(it){return it.en;}).filter(function(en){return !st.deleted.has(en);}));
    st.added.forEach(function(en){ if(enInc.has(en)){} if(itemsAll.some(function(it){return it.en===en;}) && !st.deleted.has(en)) enInc.add(en); });
    // 合并库内新增（无数值，仅管理列表显示）
    var addedList=[]; for(var ak in st.addedItems){ var ai=st.addedItems[ak]; if(!st.deleted.has(ak)) addedList.push(ai); }
    var shown=cand.filter(function(it){return enInc.has(it.en);});
    var ordered=NV.sortByCatValueDesc(shown, function(it){return it.maxVal;});
    // 热力图（行=物质，列=组别；先按行归一化，蓝色系活力分色）
    NV.heatmap(document.getElementById('viz-hm-'+kind), {rows:ordered, cols:selGroups, items:ordered, metric:metric});
    var subH=document.getElementById('viz-hm-sub-'+kind); if(subH) subH.textContent=ordered.length+' 物质 × '+selGroups.length+' 组 · 归一化'+(metric==='conc'?'浓度':'');
    var legH=document.getElementById('viz-hm-leg-'+kind); if(legH) legH.innerHTML='<span><i style="background:'+NV.bp(1)+'"></i>高</span><span><i style="background:'+NV.bp(0.5)+'"></i>中</span><span><i style="background:'+NV.bp(0)+'"></i>低</span>';
    // 类别 / 组别 chips（去重后统一，每个 chip 显示「类别 (n)」）
    NV.renderChips(kind, 'cat', [...st.cats], itemsAll.map(function(it){return it.cat||'其他';}));
    NV.renderChips(kind, 'grp', selGroups.slice().sort(), st._groupsAvail);
    // 管理列表（全部解析物质 + 库内新增，按类别·数值降序）
    var mgrAll=NV.sortByCatValueDesc(itemsAll.concat(addedList), function(it){return it.maxVal;});
    NV.renderMgr(kind, mgrAll, enInc, st);
    // 主对比矩阵与可视化面板共享删除状态：面板变动后同步刷新主表
    if(kind==='single' && window.gcRefreshMatrix) window.gcRefreshMatrix('single');
    else if(kind==='multi' && window.mcRefreshMatrix) window.mcRefreshMatrix();
  };

  NV.renderChips=function(kind, type, selected, avail){
    var box=document.getElementById('viz-'+type+'-'+kind); if(!box) return;
    // 去重 + 计数
    var counts={}; avail.forEach(function(c){ counts[c]=(counts[c]||0)+1; });
    var uniq=Object.keys(counts);
    // 按类别使用 CAT 顺序（组别按字母序）
    if(type==='cat'){
      var order=(typeof CAT!=='undefined'&&CAT)?(Array.isArray(CAT)?CAT.map(function(c){return c.name||c;}):Object.keys(CAT)):[];
      var rank=function(c){var i=order.indexOf(c);return i>=0?i:999;};
      uniq.sort(function(a,b){return rank(a)-rank(b);});
    } else { uniq.sort(); }
    var set=new Set(selected);
    box.innerHTML=uniq.map(function(c){ var on=set.has(c); var n=counts[c];
      return '<span class="viz-chip'+(on?'':' off')+'" data-type="'+type+'" data-v="'+NV.esc(c)+'" title="'+NV.esc(c)+(n>1?'（共 '+n+' 个，已合并）':'')+'">'
        +(type==='cat'?'<span class="dot" style="background:'+NV.catColor(c)+'"></span>':'')
        +NV.esc(c)+' <em class="cn">'+n+'</em></span>'; }).join('');
  };

  NV.renderMgr=function(kind, list, enInc, st){
    var box=document.getElementById('viz-mgr-'+kind); if(!box) return;
    var inc=0;
    // 按类别分组（沿用 CAT 顺序），组内按数值降序，便于按类别管理物质
    var order=(typeof CAT!=='undefined'&&CAT)?(Array.isArray(CAT)?CAT.map(function(c){return c.name||c;}):Object.keys(CAT)):[];
    var rank=function(c){var i=order.indexOf(c);return i>=0?i:999;};
    var groups={}; list.forEach(function(it){ var c=it.cat||'其他'; (groups[c]=groups[c]||[]).push(it); });
    var cats=Object.keys(groups).sort(function(a,b){return rank(a)-rank(b);});
    var html='';
    cats.forEach(function(cat){
      var items=groups[cat];
      var cz=0;
      var rows=items.map(function(it){
        var checked=enInc.has(it.en) && !st.deleted.has(it.en);
        if(checked) cz++;
        var hasVal=st._groupsAvail.some(function(g){return it.vals[g]!=null&&isFinite(it.vals[g]);});
        return '<label class="viz-mrow'+(checked?'':' del')+'"><input type="checkbox" data-en="'+NV.esc(it.en)+'" '+(checked?'checked':'')+'>'
          +'<span class="mn"><b>'+NV.esc(it.cn||it.en)+'</b></span>'
          +'<span class="mc">'+NV.esc(it.cat)+'</span>'
          +'<span class="mv">'+(hasVal?NV.fmt(it.maxVal):'无数据')+'</span></label>';
      }).join('');
      var collapsed = st.mgrCollapsed && st.mgrCollapsed.has(cat);
      html+='<div class="mgr-cat'+(collapsed?' collapsed':'')+'">'
        +'<div class="mgr-cat-h" data-cat="'+NV.esc(cat)+'"><span class="mgr-chev">▾</span><span class="dot" style="background:'+NV.catColor(cat)+'"></span>'+NV.esc(cat)+' <em class="cn">'+cz+'/'+items.length+'</em></div>'
        +'<div class="mgr-cat-body">'+rows+'</div></div>';
    });
    box.innerHTML=html;
    var inc2=0; list.forEach(function(it){ if(enInc.has(it.en) && !st.deleted.has(it.en)) inc2++; });
    inc=inc2;
    var stat=document.getElementById('viz-mgr-stat-'+kind); if(stat) stat.textContent='已纳入 '+inc+' / 共 '+list.length;
  };

  NV.wire=function(kind){
    var st=NV.state[kind];
    var $ = function(id){ return document.getElementById(id); };
    var rebuild=function(){ NV.render(kind); };
    $('viz-metric-'+kind).value=st.metric;
    $('viz-metric-'+kind).addEventListener('change', function(e){ st.metric=e.target.value; NV.save(kind); NV.render(kind); });
    var psel=$('viz-palette-'+kind);
    if(psel){ psel.value=NV._palette; psel.addEventListener('change', function(e){ NV._palette=e.target.value; try{ localStorage.setItem('gcms_nv_palette', NV._palette); }catch(_){} NV.render(kind); }); }
    var epng=$('viz-png-'+kind); if(epng) epng.addEventListener('click', function(){ NV.exportFig(kind,'png'); });
    var esvg=$('viz-svg-'+kind); if(esvg) esvg.addEventListener('click', function(){ NV.exportFig(kind,'svg'); });
    $('viz-search-'+kind).addEventListener('input', function(e){ st.search=e.target.value; NV.render(kind); });
    $('viz-save-'+kind).addEventListener('click', function(){ NV.save(kind); var b=$('viz-save-'+kind); b.textContent='已保存✓'; setTimeout(function(){b.textContent='保存调整';},1200); });
    $('viz-undo-'+kind).addEventListener('click', function(){ var en=st.undoStack.pop(); if(en){ st.deleted.delete(en); NV.save(kind); NV.render(kind); } });
    $('viz-reset-'+kind).addEventListener('click', function(){ st.deleted.clear(); st.added.clear(); st.addedItems={}; st.cats=new Set(NV.state[kind]._groupsAvail? [] : []);
      var itemsAll=(kind==='single')?NV.singleItems(st._ds,st.metric):NV.multiItems(st._ds,st.metric);
      st.cats=new Set([...new Set(itemsAll.map(function(it){return it.cat||'其他';}))]); st.groups=new Set(st._groupsAvail); NV.save(kind); NV.render(kind); });
    $('viz-all-'+kind).addEventListener('click', function(){ /* 全选=清除删除 */ st.deleted.clear(); NV.save(kind); NV.render(kind); });
    $('viz-none-'+kind).addEventListener('click', function(){ /* 全不选=全部删除 */ var itemsAll=(kind==='single')?NV.singleItems(st._ds,st.metric):NV.multiItems(st._ds,st.metric); itemsAll.forEach(function(it){ if(!st.deleted.has(it.en)){ st.deleted.add(it.en); st.undoStack.push(it.en); } }); NV.save(kind); NV.render(kind); });
    $('viz-add-'+kind).addEventListener('click', function(){ var v=$('viz-search-'+kind).value.trim(); if(!v) return; NV.addFromDB(kind, v); });
    // chips（事件委托）
    ['cat','grp'].forEach(function(type){
      var box=$('viz-'+type+'-'+kind);
      box.addEventListener('click', function(e){ var t=e.target.closest('.viz-chip'); if(!t) return;
        var v=t.getAttribute('data-v'); var set=(type==='cat')?st.cats:st.groups;
        if(set.has(v)) set.delete(v); else set.add(v); NV.save(kind); NV.render(kind); });
    });
    // 管理列表（委托）
    var mgr=$('viz-mgr-'+kind);
    mgr.addEventListener('change', function(e){ var cb=e.target; if(cb&&cb.type==='checkbox'){ var en=cb.getAttribute('data-en');
      if(cb.checked){ st.deleted.delete(en); } else { if(!st.deleted.has(en)){ st.deleted.add(en); st.undoStack.push(en); } } NV.save(kind); NV.render(kind); } });
    // 类别分组折叠 / 展开（每类可独立收起，便于长列表管理）
    mgr.addEventListener('click', function(e){ var h=e.target.closest('.mgr-cat-h'); if(!h) return;
      var cat=h.getAttribute('data-cat'); var catDiv=h.closest('.mgr-cat'); if(!catDiv) return;
      var set=st.mgrCollapsed || (st.mgrCollapsed=new Set());
      if(set.has(cat)) set.delete(cat); else set.add(cat);
      catDiv.classList.toggle('collapsed'); });
  };

  NV.addFromDB=function(kind, name){
    var st=NV.state[kind];
    var finish=function(db){
      if(!db){ alert('未能读取化合物库'); return; }
      var n=name.trim().toLowerCase();
      var hit=db.find(function(c){ return (c.en&&c.en.toLowerCase()===n)||(c.cn&&c.cn.toLowerCase()===n)||((c.syn||[]).some&&(c.syn||[]).some(function(s){return s.toLowerCase()===n;})); });
      if(!hit){ alert('库内未找到：'+name); return; }
      var en=hit.en; st.added.add(en); st.deleted.delete(en);
      st.addedItems[en]={en:en, cn:(hit.cn||hit.en), cat:(hit.cat||'其他'), odor:(hit.odor||''), vals:{}, maxVal:0};
      NV.save(kind); NV.render(kind);
    };
    if(typeof DB_RAW!=='undefined' && DB_RAW) finish(DB_RAW);
    else if(typeof _allDB!=='undefined' && _allDB) finish(_allDB);
    else { fetch('/api/all').then(function(r){return r.json();}).then(function(d){ _allDB=d; finish(d); }).catch(function(){ alert('未能读取化合物库'); }); }
  };

  NV.state={};
  window.NatureViz=NV;
})();
/* ===== NATURE_VIZ_END ===== */
</script>
</body>
</html>
"""

import json
@app.route('/')
def index():
    return render_template_string(PAGE, cat_json=json.dumps(CAT_COLOR, ensure_ascii=False))

@app.route('/api/info')
def api_info():
    return jsonify({'size': len(COMPOUNDS)})

@app.route('/api/all')
def api_all():
    """返回全部风味化合物（用于浏览数据库），仅暴露展示字段。
    排序：优先展示「有风味描述 + 有阈值 + 有来源」的物质（信息完整者在前）。"""
    fields = ('en', 'cn', 'cas', 'cat', 'thr', 'med', 'odor', 'source', 'syn')
    out = [{k: c.get(k, '') for k in fields} for c in COMPOUNDS]
    out = sort_by_completeness(out)
    return jsonify(out)

@app.route('/api/search', methods=['POST'])
def api_search():
    data = request.get_json(force=True)
    names = data.get('names', [])
    results = [match_compound(n) for n in names]
    # 命中多个候选时，信息完整（有风味描述/阈值/来源）的优先展示
    results = sort_by_completeness(results)
    return jsonify(results)

def _read_text(path):
    """读取文本/Markdown 文件内容（utf-8 优先，失败回退 gbk/latin）。"""
    for enc in ('utf-8', 'utf-8-sig', 'gbk', 'latin-1'):
        try:
            return open(path, 'r', encoding=enc).read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    return ''


def _to_float(x):
    if x is None:
        return None
    s = str(x).replace(',', '').strip()
    if s in ('', 'ND', 'nd', 'N.D.', 'N/A', 'na', '—', '-', '–', 'NULL', 'null'):
        return None
    try:
        return float(s)
    except (ValueError, TypeError):
        return None


def _parse_pdf(path, ist_name='2-辛醇'):
    """用 pdfplumber 解析 PDF：优先结构化表格（Agilent MassHunter 定量报告），
    否则回退到文本扫描。返回带 en/rt/resp/rr/conc0/istd_area 的行列表（供 enrich）。"""
    ist_name = (ist_name or '2-辛醇').strip() or '2-辛醇'
    # 优先尝试多处理组 Shimadzu 定量报告（一次解析自动识别各处理组数据）
    try:
        multi = parse_gcms_pdf_multi(path)
        if multi and multi.get('compounds'):
            if len(multi.get('groups', [])) >= 2:
                return multi
            # 单处理组：转为普通单行结构，复用既有单报告渲染
            rows = multi.get('rows') or []
            ist_area = next(iter((multi.get('ist_area_by_group') or {}).values()), None)
            return {'rows': rows, 'ist_name': multi.get('ist_name'),
                    'ist_area': ist_area, 'schema': 'shimadzu_single'}
    except Exception as e:
        print('[warn] 多处理组解析失败，回退单报告：', e, file=sys.stderr)
    rows = _parse_pdf_tables(path, ist_name)
    if rows:
        return rows
    # 回退：纯文本扫描（兼容无清晰表格的 PDF）
    text = ''
    with pdfplumber.open(path) as pdf:
        for pg in pdf.pages:
            text += (pg.extract_text() or '') + '\n'
    if not text.strip():
        return []
    return parse_report_text(text, ist_name=ist_name)


def _parse_pdf_tables(path, ist_name='2-辛醇'):
    """结构化解析 Agilent MassHunter 定量报告表格。

    表头：化合物 | ISTD | RT | 响应 | ISTD 响应 | 响应比 | 最终浓度 | 单位
    每行首列是化合物名（中/英文，可能含括号/连字符/换行），其余列对齐稳定。
    通过表头语义定位各数值列，避免「ISTD 响应」拆词、化合物名含空格导致错位。
    返回行列表（en = 命中库标准名或原名；含 rt/resp/rr/conc0/istd_area）。"""
    ist_name = (ist_name or '2-辛醇').strip() or '2-辛醇'
    out, istd_area, seen_en = [], None, set()
    try:
        with pdfplumber.open(path) as pdf:
            for pg in pdf.pages:
                for t in pg.extract_tables():
                    # 定位表头行（含「化合物」的那一行）
                    hidx = None
                    for i, r in enumerate(t):
                        if r and any((c or '').strip() == '化合物' for c in r):
                            hidx = i
                            break
                    if hidx is None:
                        continue
                    header = [(c or '').strip() for c in t[hidx]]

                    def _col(*names):
                        for n in names:
                            for j, h in enumerate(header):
                                if h == n:
                                    return j
                        return None
                    ci = {
                        'comp': header.index('化合物') if '化合物' in header else 0,
                        'rt':   _col('RT', '保留时间', '出峰时间'),
                        'resp': _col('响应', '峰面积', '响应值'),
                        'istd': _col('ISTD 响应', '内标峰面积', '内标响应'),
                        'rr':   _col('响应比', '相对响应', '相对峰面积'),
                        'conc': _col('最终浓度', '浓度', '含量'),
                    }
                    for r in t[hidx + 1:]:
                        if not r:
                            continue
                        comp = (r[ci['comp']] or '').strip() if ci['comp'] < len(r) else ''
                        if not comp:
                            continue
                        if comp in ('定量分析完成报告', '样品类型', '化合物', '样品色谱图', 'RT'):
                            continue
                        rt = _to_float(r[ci['rt']]) if ci['rt'] is not None and ci['rt'] < len(r) else None
                        resp = _to_float(r[ci['resp']]) if ci['resp'] is not None and ci['resp'] < len(r) else None
                        istd = _to_float(r[ci['istd']]) if ci['istd'] is not None and ci['istd'] < len(r) else None
                        rr = _to_float(r[ci['rr']]) if ci['rr'] is not None and ci['rr'] < len(r) else None
                        conc_cell = r[ci['conc']] if ci['conc'] is not None and ci['conc'] < len(r) else None
                        conc0 = _to_float(conc_cell)
                        if istd_area is None and istd is not None:
                            istd_area = istd
                        m = match_compound(comp)
                        en = m.get('en') or comp
                        if en in seen_en:   # 同报告内按标准名去重，避免误匹配导致浓度重复累计
                            continue
                        seen_en.add(en)
                        item = {
                            'en': en,
                            'rt': rt, 'resp': resp, 'rr': rr,
                            'conc0': conc0, 'istd_area': istd_area,
                            'cn': m.get('cn'), 'cat': m.get('cat', '其他'),
                            'thr': m.get('thr'), 'note': m.get('note'),
                            'match': m.get('match'),
                        }
                        out.append(item)
    except Exception as e:
        print('[warn] PDF 表格解析失败，回退文本扫描：', e, file=sys.stderr)
        return {'rows': [], 'ist_name': ist_name, 'ist_area': None,
                'schema': 'pdf', 'headers': []}
    return {'rows': out, 'ist_name': ist_name, 'ist_area': istd_area,
            'schema': 'pdf', 'headers': header}


def _read_docx(path):
    """用 python-docx 提取正文文本。"""
    import docx
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # 表格单元格也纳入
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paras.append(cell.text)
    return '\n'.join(paras)


def _parse_batch_xlsx(path):
    """读取 xlsx 首个工作表，提取所有非空文本单元格作为候选，再识别库内物质。"""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb[wb.sheetnames[0]]
    cells = []
    for r in ws.iter_rows(values_only=True):
        for v in r:
            if isinstance(v, str) and v.strip():
                cells.append(v.strip())
            elif isinstance(v, (int, float)):
                cells.append(str(v))
    wb.close()
    if not cells:
        return []
    text = '\n'.join(cells)
    # 先逐单元格精确匹配（如单元格就是物质名），再整体扫描兜底
    names = extract_substance_names(text)
    return names


def _parse_batch_csv(path):
    """读取 csv，合并所有单元格文本后扫描识别库内物质。"""
    import csv
    rows = []
    with open(path, 'r', encoding='utf-8-sig', errors='ignore') as f:
        for r in csv.reader(f):
            rows.extend([c.strip() for c in r if c and c.strip()])
    if not rows:
        return []
    return extract_substance_names('\n'.join(rows))


@app.route('/api/batch_upload', methods=['POST'])
def api_batch_upload():
    """批量检索文件上传：支持 xlsx / pdf / csv / docx。
    从文件中识别库内已知风味物质，返回标准英文名列表（去重、保序）。"""
    f = request.files.get('file')
    if not f:
        return jsonify({'names': []})
    suffix = (os.path.splitext(f.filename)[-1] or '.txt').lower()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        f.save(tmp.name)
        path = tmp.name
    try:
        if suffix in ('.xlsx', '.xls'):
            names = _parse_batch_xlsx(path)
        elif suffix in ('.pdf',):
            names = extract_substance_names(_parse_pdf_text(path))
        elif suffix in ('.docx',):
            names = extract_substance_names(_read_docx(path))
        elif suffix in ('.csv',):
            names = _parse_batch_csv(path)
        else:
            # 文本类：直接扫描
            names = extract_substance_names(_read_text(path))
        return jsonify({'names': names, 'count': len(names)})
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass


def _parse_pdf_text(path):
    """仅提取 PDF 文本（供批量扫描复用）。"""
    text = ''
    with pdfplumber.open(path) as pdf:
        for pg in pdf.pages:
            text += (pg.extract_text() or '') + '\n'
    return text




@app.route('/api/upload', methods=['POST'])
def api_upload():
    f = request.files.get('file')
    if not f:
        return jsonify([])
    suffix = (os.path.splitext(f.filename)[-1] or '.xlsx').lower()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        f.save(tmp.name)
        path = tmp.name
    try:
        ist_name = request.form.get('ist_name', '2-辛醇').strip() or '2-辛醇'
        if suffix in ('.pdf',):
            data = _parse_pdf(path, ist_name=ist_name)
        elif suffix in ('.docx',):
            data = parse_report_text(_read_docx(path), ist_name=ist_name)
        elif suffix in ('.md', '.markdown', '.txt', '.csv'):
            data = parse_report_text(_read_text(path), ist_name=ist_name)
        else:
            data = parse_gcms_excel(path)
        # 多处理组报告：直接透传（已含 en/cn/cat/thr 与各组 by_group 数据）
        if isinstance(data, dict) and data.get('multi'):
            return jsonify(data)
        if not isinstance(data, dict):
            data = {'rows': data or []}
        rows = data.get('rows') or []
        if not rows:
            return jsonify({'rows': [], 'ist_name': data.get('ist_name'),
                            'ist_area': data.get('ist_area'), 'schema': data.get('schema')})
        results = enrich(rows)
        # 排序：先按类别分组（同类聚在一起），组内按响应比降序（气味活性最高在前）
        cat_order = list(CAT_COLOR.keys()) + ['未匹配', '其他', '']
        def _cat_idx(c):
            c = (c or '').strip()
            return cat_order.index(c) if c in cat_order else len(cat_order)
        results.sort(key=lambda r: (_cat_idx(r.get('cat')),
                                     -(r.get('rr') if isinstance(r.get('rr'), (int, float)) else -1)))
        return jsonify({'rows': results, 'ist_name': data.get('ist_name'),
                        'ist_area': data.get('ist_area'), 'schema': data.get('schema')})
    except Exception as e:
        # 解析/富集中任何异常都返回 JSON 错误（而非 HTML 500），避免前端静默崩溃
        return jsonify({'error': '报告解析失败：' + str(e), 'rows': []})
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass

@app.route('/api/multi_upload', methods=['POST'])
def api_multi_upload():
    """多报告批量对比：接收多个文件(字段名 files[]) 与对应样品名(字段名 names[])。
    每个文件独立解析+enrich，返回 [{name, rows:[enrich结果]}]。浓度/OAV 由前端按统一内标参数计算。"""
    files = request.files.getlist('files')
    names = request.form.getlist('names')
    ist_name = (request.form.get('ist_name', '2-辛醇') or '2-辛醇').strip() or '2-辛醇'
    if not files:
        return jsonify({'samples': []})
    samples = []
    try:
        for idx, f in enumerate(files):
            suffix = (os.path.splitext(f.filename)[-1] or '.xlsx').lower()
            name = (names[idx] if idx < len(names) and names[idx].strip() else None) or f.filename or f'样品{idx+1}'
            # 去扩展名
            name = re.sub(r'\.[^.]+$', '', str(name))
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                f.save(tmp.name)
                path = tmp.name
            try:
                if suffix in ('.pdf',):
                    data = _parse_pdf(path, ist_name=ist_name)
                elif suffix in ('.docx',):
                    data = parse_report_text(_read_docx(path), ist_name=ist_name)
                elif suffix in ('.md', '.markdown', '.txt'):
                    data = parse_report_text(_read_text(path), ist_name=ist_name)
                else:
                    data = parse_gcms_excel(path)
                # 多处理组 PDF：把每个处理组展开为一个样品，复用多报告对比
                if isinstance(data, dict) and data.get('multi') and len(data.get('groups', [])) >= 2:
                    for gi, g in enumerate(data.get('groups', [])):
                        grows = [r for r in data.get('rows', []) if r.get('group') == g]
                        gname = (names[idx] if idx < len(names) and names[idx].strip() else '') or f'{name}·{g}'
                        gname = re.sub(r'\.[^.]+$', '', str(gname))
                        samples.append({
                            'name': gname, 'rows': enrich(grows),
                            'ist_name': data.get('ist_name'),
                            'ist_area': data.get('ist_area_by_group', {}).get(g),
                        })
                    continue
                if not isinstance(data, dict):
                    data = {'rows': data or []}
                rows = data.get('rows') or []
                if not rows:
                    samples.append({'name': name, 'rows': [], 'empty': True,
                                    'ist_name': data.get('ist_name'),
                                    'ist_area': data.get('ist_area')})
                    continue
                results = enrich(rows)
                samples.append({'name': name, 'rows': results,
                               'ist_name': data.get('ist_name'),
                               'ist_area': data.get('ist_area')})
            finally:
                try:
                    os.unlink(path)
                except Exception:
                    pass
    except Exception as e:
        return jsonify({'samples': [], 'error': '多报告解析失败：' + str(e)})
    return jsonify({'samples': samples})

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port, debug=False)
